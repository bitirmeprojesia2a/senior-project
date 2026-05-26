"""Transcript extraction and compact question answering helpers.

This module intentionally keeps OCR/text extraction outside the main RAG flow.
Uploaded transcript facts are user-scoped context, not institutional sources.
"""

from __future__ import annotations

import io
import json
import re
import tempfile
import time
from dataclasses import dataclass, field, replace
from datetime import time as clock_time
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any, Protocol

from src.core.text_normalization import contains_any_normalized, normalize_text
from src.quality.answer_filter import REWRITE_ONLY_SYSTEM_SUFFIX, check_answer_quality, quality_issue_blocks_answer
from src.transcripts.document_intent import DocumentQuestionIntent, classify_document_question


_MAX_STORED_COURSES = 240
_MAX_REASONABLE_TRANSCRIPT_AKTS = 450
_MAX_REASONABLE_TRANSCRIPT_CREDIT = 300
_MAX_REASONABLE_COURSE_AKTS = 60
_MAX_REASONABLE_COURSE_CREDIT = 30
_DEFAULT_TTL_SECONDS = 60 * 60 * 6
_COURSE_CODE_RE = re.compile(r"\b([A-ZÇĞİÖŞÜ]{2,8}\s?\d{3,4}[A-Z]?)\b", re.IGNORECASE)
_GRADE_RE = re.compile(r"\b(AA|BA|BB|CB|CC|DC|DD|FD|FF|YT|YZ|G|K|P|F)\b", re.IGNORECASE)
_NUMBER_RE = re.compile(r"(?<!\d)(\d{1,3}(?:[,.]\d{1,2})?)(?!\d)")
_FAILED_GRADES = {"FF", "FD", "K", "YZ", "F"}
_PASSED_GRADES = {"AA", "BA", "BB", "CB", "CC", "DC", "DD", "YT", "G", "P"}
_DETAILED_GRADES = {"AA", "BA", "BB", "CB", "CC", "DC", "DD", "FD", "FF", "YT", "YZ"}
_TEXT_EXTENSIONS = {".txt", ".csv"}
_PDF_EXTENSIONS = {".pdf"}
_DOCX_EXTENSIONS = {".docx"}
_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp"}
_MIN_PDF_TEXT_CHARS = 80
_OCR_PDF_MAX_PAGES = 8
_OCR_PDF_DPI = 220
_EXPLICIT_TRANSCRIPT_MARKERS = (
    "transkript",
    "not dokumu",
    "not dokum",
)
_PERSONAL_TRANSCRIPT_MARKERS = (
    "aktsim",
    "kredim",
    "gectigim",
    "kaldigim",
    "aldigim",
    "aldigi",
    "aldiklarim",
    "aldiklari",
    "almis",
    "almisim",
    "aldim",
    "dersim",
    "derslerim",
    "notum",
    "notlarim",
    "ortalamam",
)
_TRANSCRIPT_METRIC_MARKERS = (
    "toplam akts",
    "toplam kredi",
    "kac akts",
    "kac kredi",
    "basarisiz",
    "basarili ders",
    "gno",
    "agno",
    "not",
    "harf notu",
)
_INSTITUTIONAL_AKTS_MARKERS = (
    "mezun",
    "mezuniyet",
    "program",
    "lisans",
    "onlisans",
    "dis hekimligi",
    "tip",
    "muhendisligi",
    "ogretmenligi",
    "icin kac akts",
)

_FIELD_SEPARATOR_RE = re.compile(r"^\s*([^:：|]{2,55})\s*[:：|]\s*(.{1,180})\s*$")
_FIELD_SPACED_RE = re.compile(r"^\s*(.{2,55}?)\s{2,}(.{1,180})\s*$")
_FIELD_KEY_CLEAN_RE = re.compile(r"[^a-z0-9]+")
_KNOWN_FIELD_LABELS = (
    "ogrencilik durumu",
    "ogrenci numarasi",
    "ogrenci no",
    "tc kimlik numarasi",
    "tc kimlik no",
    "t c kimlik no",
    "t c kimlik numarasi",
    "kimlik no",
    "adi soyadi",
    "adi ve soyadi",
    "ad soyad",
    "ad soyadi",
    "ad soyadi",
    "e posta adresi",
    "eposta adresi",
    "e mail",
    "email",
    "telefon no",
    "telefon numarasi",
    "ogretim yili",
    "web adresi",
    "faks no",
    "gorev ve unvani",
    "imza kase",
    "imza ve kase",
    "baba adi",
    "anne adi",
    "dogum tarihi",
    "dogum yeri",
    "belge tarihi",
    "barkod numarasi",
    "barkod no",
    "kayit tarihi",
    "giris tarihi",
    "mezuniyet tarihi",
    "program adi",
    "program",
    "bolum adi",
    "bolum",
    "fakulte",
    "sinif",
    "sinifi",
    "durumu",
    "uyruk",
    "kan grubu",
)
_FIELD_LABEL_ALIASES = (
    ("tc kimlik no", ("tc kimlik", "tc kimlik no", "kimlik no", "tc kimlik numarasi", "t c kimlik no", "t c kimlik numarasi")),
    ("adi soyadi", ("adi soyadi", "adi ve soyadi", "ad soyad", "ad soyadi", "kisinin adi", "adi kisinin")),
    ("e posta adresi", ("e posta", "eposta", "e posta adresi", "eposta adresi", "email", "e mail")),
    ("telefon no", ("telefon", "telefon no", "telefon numarasi")),
)
_FIELD_KEY_SKIP_MARKERS = (
    "http",
    "www",
    "pdf",
    "sayfa",
    "ondokuz mayis universitesi",
    "yok transkript",
)


class TranscriptProcessingError(RuntimeError):
    """Raised when a transcript file cannot be extracted safely."""


class LLMGenerateProtocol(Protocol):
    async def generate(
        self,
        prompt: str,
        system: str | None = None,
        json_mode: bool = False,
        *,
        model: str | None = None,
        model_role: str = "default",
        llm_profile: str | None = None,
    ) -> str:
        ...


@dataclass(frozen=True)
class TranscriptCourse:
    code: str
    name: str
    credit: float | None = None
    akts: float | None = None
    grade: str | None = None
    status: str = "unknown"
    term: str | None = None


@dataclass(frozen=True)
class DocumentField:
    key: str
    label: str
    value: str
    confidence: str = "medium"
    state: str = "filled"
    source: str = "text_layer"


@dataclass(frozen=True)
class DocumentFacts:
    """Common fact layer for Slack uploaded documents.

    This is intentionally scoped to uploaded documents. Institutional RAG
    ingestion keeps using its own source/index pipeline.
    """

    document_type: str
    title: str | None = None
    issuer: str | None = None
    person_name: str | None = None
    fields: tuple[DocumentField, ...] = ()
    tables: tuple[dict[str, Any], ...] = ()
    schedule_slots: tuple[dict[str, Any], ...] = ()
    transcript_courses: tuple[TranscriptCourse, ...] = ()
    dates: tuple[DocumentField, ...] = ()
    status: str | None = None
    extraction_confidence: str = "medium"
    warnings: tuple[str, ...] = ()
    raw_text_preview: str = ""


@dataclass(frozen=True)
class _DocumentFieldMatch:
    field: DocumentField
    score: float


@dataclass(frozen=True)
class _TermRequest:
    academic_year: str | None
    season: str | None
    display: str
    single_year: int | None = None


@dataclass(frozen=True)
class TranscriptDocument:
    filename: str
    text: str
    document_type: str = "transcript"
    facts: DocumentFacts | None = None
    fields: tuple[DocumentField, ...] = ()
    schedule_slots: tuple[dict[str, Any], ...] = ()
    courses: tuple[TranscriptCourse, ...] = ()
    total_akts: float | None = None
    total_credit: float | None = None
    gpa: float | None = None
    program_name: str | None = None
    parse_confidence: str = "medium"
    extraction_mode: str = "text"
    warnings: tuple[str, ...] = ()
    extraction_debug: dict[str, Any] = field(default_factory=dict)

    @property
    def passed_courses(self) -> tuple[TranscriptCourse, ...]:
        return tuple(course for course in self.courses if course.status == "passed")

    @property
    def failed_courses(self) -> tuple[TranscriptCourse, ...]:
        return tuple(course for course in self.courses if course.status == "failed")


@dataclass
class _StoredTranscript:
    document: TranscriptDocument
    user_id: str
    last_query: str | None = None
    stored_at: float = field(default_factory=time.monotonic)


class TranscriptSessionStore:
    """Small TTL store keyed by Slack conversation context.

    Keeping this separate from ConversationContextService prevents uploaded
    personal documents from becoming normal RAG/conversation memory.
    """

    def __init__(self, *, ttl_seconds: int = _DEFAULT_TTL_SECONDS) -> None:
        self.ttl_seconds = ttl_seconds
        self._items: dict[str, _StoredTranscript] = {}

    def put(
        self,
        *,
        context_id: str,
        user_id: str,
        document: TranscriptDocument,
        last_query: str | None = None,
    ) -> None:
        self._purge_expired()
        self._items[context_id] = _StoredTranscript(document=document, user_id=user_id, last_query=last_query)

    def get(self, *, context_id: str, user_id: str) -> TranscriptDocument | None:
        self._purge_expired()
        item = self._items.get(context_id)
        if item is None or item.user_id != user_id:
            return None
        return item.document

    def get_last_query(self, *, context_id: str, user_id: str) -> str | None:
        self._purge_expired()
        item = self._items.get(context_id)
        if item is None or item.user_id != user_id:
            return None
        return item.last_query

    def update_last_query(self, *, context_id: str, user_id: str, last_query: str | None) -> None:
        self._purge_expired()
        item = self._items.get(context_id)
        if item is None or item.user_id != user_id:
            return
        item.last_query = last_query
        item.stored_at = time.monotonic()

    def _purge_expired(self) -> None:
        now = time.monotonic()
        expired = [
            key
            for key, value in self._items.items()
            if now - value.stored_at > self.ttl_seconds
        ]
        for key in expired:
            self._items.pop(key, None)


class TranscriptProcessor:
    """Extract uploaded document text/facts and answer scoped questions."""

    def __init__(self, *, llm_service: LLMGenerateProtocol | None = None) -> None:
        self.llm_service = llm_service

    def process_bytes(self, *, filename: str, content: bytes, mimetype: str | None = None) -> TranscriptDocument:
        if not content:
            raise TranscriptProcessingError("Dosya boş görünüyor.")
        extension = Path(filename or "").suffix.lower()
        pre_extracted_form_fields: tuple[DocumentField, ...] = ()
        try:
            text, mode, warnings = self._extract_text(
                filename=filename,
                content=content,
                mimetype=mimetype,
                extension=extension,
            )
        except TranscriptProcessingError:
            if extension not in _PDF_EXTENSIONS:
                raise
            pre_extracted_form_fields = tuple(_extract_pdf_form_fields(content))
            if not pre_extracted_form_fields:
                raise
            warnings = ["PDF metin/OCR cikaramadi; PDF form alanlari kullanildi."]
            mode = "pdf_form_fields"
            text = "\n".join(
                f"{field.label}: {field.value}" if field.value.strip() else field.label
                for field in pre_extracted_form_fields
            )
        if not text.strip():
            raise TranscriptProcessingError("Dosyadan okunabilir transkript metni çıkarılamadı.")

        document_type = _classify_document(filename=filename, text=text)
        if document_type != "transcript":
            text_fields = tuple(_extract_document_fields(text))
            form_fields = (
                pre_extracted_form_fields
                if pre_extracted_form_fields
                else tuple(_extract_pdf_form_fields(content)) if extension in _PDF_EXTENSIONS else ()
            )
            fields = _merge_document_fields(form_fields, text_fields)
            schedule_slots: tuple[dict[str, Any], ...] = ()
            if document_type == "schedule_document" and extension in _PDF_EXTENSIONS:
                schedule_slots = tuple(_extract_schedule_slots_from_pdf_content(filename=filename, content=content))
                if not schedule_slots:
                    schedule_slots = tuple(_extract_schedule_slots_from_text(filename=filename, text=text))
                if schedule_slots:
                    warnings.append(f"Ders programı tablosundan {len(schedule_slots)} satır yapılandırıldı.")
            parse_confidence = "high" if schedule_slots else ("medium" if len(text.strip()) >= 160 else "low")
            facts = _build_document_facts(
                filename=filename or "belge",
                text=text,
                document_type=document_type,
                fields=fields,
                schedule_slots=schedule_slots,
                courses=(),
                total_akts=None,
                total_credit=None,
                gpa=None,
                program_name=None,
                extraction_confidence=parse_confidence,
                warnings=tuple(warnings),
            )
            return TranscriptDocument(
                filename=filename or "belge",
                text=_compact_raw_text(text),
                document_type=document_type,
                facts=facts,
                fields=fields,
                schedule_slots=schedule_slots,
                extraction_mode=mode,
                warnings=tuple(warnings),
                parse_confidence=parse_confidence,
                extraction_debug=_build_extraction_debug(
                    text=text,
                    document_type=document_type,
                    extraction_mode=mode,
                    warnings=warnings,
                    fields=fields,
                    courses=(),
                    parse_confidence=parse_confidence,
                ),
            )

        courses = tuple(_parse_courses(text))[:_MAX_STORED_COURSES]
        total_akts = _extract_total_value(text, kind="akts")
        total_credit = _extract_total_value(text, kind="credit")
        if total_akts is None and _has_unreasonable_total_literal(text, kind="akts"):
            warnings.append("Toplam AKTS degeri guvenilir gorunmedigi icin gosterilmedi.")
        if total_credit is None and _has_unreasonable_total_literal(text, kind="credit"):
            warnings.append("Toplam kredi degeri guvenilir gorunmedigi icin gosterilmedi.")
        if not _is_reasonable_total(total_akts, kind="akts"):
            if total_akts is not None:
                warnings.append("Toplam AKTS degeri guvenilir gorunmedigi icin gosterilmedi.")
            total_akts = None
        if not _is_reasonable_total(total_credit, kind="credit"):
            if total_credit is not None:
                warnings.append("Toplam kredi degeri guvenilir gorunmedigi icin gosterilmedi.")
            total_credit = None
        if total_akts is None:
            summed_akts = _sum_reliable_course_metric(courses, metric="akts")
            total_akts = summed_akts if summed_akts is not None and summed_akts > 0 else None
        if total_credit is None:
            summed_credit = _sum_reliable_course_metric(courses, metric="credit")
            total_credit = summed_credit if summed_credit is not None and summed_credit > 0 else None

        parse_confidence = _transcript_parse_confidence(
            courses=courses,
            total_akts=total_akts,
            total_credit=total_credit,
        )
        gpa = _extract_gpa(text)
        program_name = _extract_program_name(text)
        facts = _build_document_facts(
            filename=filename or "transkript",
            text=text,
            document_type=document_type,
            fields=(),
            schedule_slots=(),
            courses=courses,
            total_akts=total_akts,
            total_credit=total_credit,
            gpa=gpa,
            program_name=program_name,
            extraction_confidence=parse_confidence,
            warnings=tuple(warnings),
        )
        return TranscriptDocument(
            filename=filename or "transkript",
            text=_compact_raw_text(text),
            document_type=document_type,
            facts=facts,
            courses=courses,
            total_akts=total_akts,
            total_credit=total_credit,
            gpa=gpa,
            program_name=program_name,
            parse_confidence=parse_confidence,
            extraction_mode=mode,
            warnings=tuple(warnings),
            extraction_debug=_build_extraction_debug(
                text=text,
                document_type=document_type,
                extraction_mode=mode,
                warnings=warnings,
                fields=(),
                courses=courses,
                parse_confidence=parse_confidence,
            ),
        )

    async def answer_question(
        self,
        *,
        document: TranscriptDocument,
        query: str,
        llm_profile: str | None = None,
        document_intent: DocumentQuestionIntent | str | None = None,
    ) -> str:
        intent = _coerce_document_intent(document_intent)
        if document.document_type != "transcript":
            schedule_answer = _answer_schedule_document_question(document=document, query=query, document_intent=intent)
            if schedule_answer is not None:
                return schedule_answer
            return await self._answer_generic_document_question(
                document=document,
                query=query,
                llm_profile=llm_profile,
                document_intent=intent,
            )
        deterministic = _answer_transcript_question(document=document, query=query)
        if deterministic is not None:
            return deterministic
        if self.llm_service is None:
            return _format_transcript_summary(document)

        prompt = _build_llm_transcript_prompt(document=document, query=query)
        try:
            answer = await self.llm_service.generate(
                prompt,
                system=_TRANSCRIPT_QA_SYSTEM_PROMPT,
                model_role="final_refinement",
                llm_profile=llm_profile,
            )
        except Exception:
            return _format_transcript_summary(document)
        cleaned = (answer or "").strip()
        return cleaned or _format_transcript_summary(document)

    async def _answer_generic_document_question(
        self,
        *,
        document: TranscriptDocument,
        query: str,
        llm_profile: str | None,
        document_intent: DocumentQuestionIntent | None = None,
    ) -> str:
        semantic_answer = _answer_generic_document_semantic_question(document=document, query=query)
        if semantic_answer is not None:
            return semantic_answer
        classification = None if document_intent else classify_document_question(query=query, document=document)
        effective_intent = document_intent or (
            classification.intent if classification is not None else DocumentQuestionIntent.DOCUMENT_SUMMARY
        )
        if effective_intent == DocumentQuestionIntent.DOCUMENT_FIELD_LOOKUP:
            match = _find_document_field_match(document, query)
            field = match.field if match is not None else None
            if self.llm_service is not None and (field is None or match is None or match.score < 0.88):
                mapped_field = await _map_document_field_with_llm(
                    llm_service=self.llm_service,
                    document=document,
                    query=query,
                    llm_profile=llm_profile,
                )
                if mapped_field is not None:
                    field = mapped_field
            if field is not None:
                if field.state != "filled" or not field.value.strip():
                    return (
                        f"Belgede {field.label} alani var; ancak bu alana ait deger "
                        "doldurulmus ya da guvenilir bicimde okunmus gorunmuyor."
                    )
                return f"Belgede {field.label} {field.value} olarak gorunuyor."
            requested_field = _requested_field_label(query)
            label = f"\"{requested_field}\" alanini" if requested_field else "istenen alani"
            return f"Bu belgede {label} guvenilir bicimde bulamadim."
        if self.llm_service is None:
            return _format_generic_document_summary(document)
        prompt = _build_llm_document_prompt(document=document, query=query, document_intent=effective_intent)
        try:
            answer = await self.llm_service.generate(
                prompt,
                system=_GENERIC_DOCUMENT_QA_SYSTEM_PROMPT,
                model_role="final_refinement",
                llm_profile=llm_profile,
            )
        except Exception:
            return _format_generic_document_summary(document)
        cleaned = (answer or "").strip()
        if cleaned:
            cleaned = await self._repair_generic_document_answer_if_needed(
                prompt=prompt,
                answer=cleaned,
                llm_profile=llm_profile,
                fallback=_format_generic_document_summary(document),
            )
        return cleaned or _format_generic_document_summary(document)

    async def _repair_generic_document_answer_if_needed(
        self,
        *,
        prompt: str,
        answer: str,
        llm_profile: str | None,
        fallback: str,
    ) -> str:
        if self.llm_service is None:
            return answer
        quality = check_answer_quality(answer)
        if not quality.needs_rewrite:
            return answer
        repair_prompt = (
            f"{prompt}\n\nOnceki cevap:\n{answer}\n\n"
            "Gorev: Onceki cevabi ayni bilgiyle, bozuk/yabanci token kullanmadan dogal Turkce yeniden yaz."
        )
        try:
            repaired = await self.llm_service.generate(
                repair_prompt,
                system=_GENERIC_DOCUMENT_QA_SYSTEM_PROMPT + REWRITE_ONLY_SYSTEM_SUFFIX,
                model_role="final_refinement",
                llm_profile=llm_profile,
            )
        except Exception:
            return fallback if quality_issue_blocks_answer(quality) else answer
        repaired_answer = (repaired or "").strip()
        repaired_quality = check_answer_quality(repaired_answer)
        if repaired_answer and not quality_issue_blocks_answer(repaired_quality):
            return repaired_answer
        return fallback if quality_issue_blocks_answer(quality) else answer

    def _extract_text(
        self,
        *,
        filename: str,
        content: bytes,
        mimetype: str | None,
        extension: str,
    ) -> tuple[str, str, list[str]]:
        warnings: list[str] = []
        normalized_mimetype = (mimetype or "").lower()
        if extension in _PDF_EXTENSIONS or "pdf" in normalized_mimetype:
            try:
                text = _extract_pdf_text(content)
            except TranscriptProcessingError:
                text = ""
                warnings.append("PDF metin katmanı okunamadı; OCR denenmiştir.")
            if len(text.strip()) >= _MIN_PDF_TEXT_CHARS:
                return text, "pdf_text", warnings
            warnings.append("PDF metin katmanı zayıf görünüyor; OCR denenmiştir.")
            ocr_text = _extract_pdf_ocr_text(content)
            if ocr_text.strip():
                return ocr_text, "pdf_ocr", warnings
            if text.strip():
                warnings.append("OCR metin çıkaramadı; PDF metin katmanındaki sınırlı içerik kullanıldı.")
                return text, "pdf_text_partial", warnings
            if _looks_like_course_schedule_document(
                normalized_name=normalize_text(filename or ""),
                normalized_text="",
            ):
                warnings.append(
                    "PDF metin/OCR çıktısı boş; ders programı olduğu dosya adından anlaşıldığı için yapılandırılmış ders programı kaynakları kullanılmalı."
                )
                return filename, "pdf_filename_fallback", warnings
            raise TranscriptProcessingError(
                "PDF transkriptten metin çıkarılamadı. Taranmış PDF için OCR motoru hazır değil veya sonuç boş."
            )
        if extension in _DOCX_EXTENSIONS:
            return _extract_docx_text(content), "docx_text", warnings
        if extension in _TEXT_EXTENSIONS or normalized_mimetype.startswith("text/"):
            return _decode_text(content), "plain_text", warnings
        if extension in _IMAGE_EXTENSIONS or normalized_mimetype.startswith("image/"):
            text = _extract_image_ocr_text(content)
            if text:
                return text, "image_ocr", warnings
            warnings.append("Görsel OCR motoru kurulu değil veya metin çıkarılamadı.")
            raise TranscriptProcessingError(
                "Görsel transkript için OCR motoru şu ortamda hazır değil. PDF veya metin tabanlı transkript yükleyin."
            )
        return _decode_text(content), "plain_text_fallback", warnings


def is_transcript_followup_query(query: str | None) -> bool:
    raw_lower = (query or "").lower()
    normalized = normalize_text(query or "")
    if _forces_uploaded_document_context(normalized):
        return True
    if re.search(r"\b(?:akts|kredi)\s*['’]?\s*(?:im|ım|m)\b", raw_lower):
        return True
    normalized = normalize_text(query or "")
    if not normalized:
        return False
    tokens = set(normalized.split())
    if any(marker in normalized for marker in _EXPLICIT_TRANSCRIPT_MARKERS):
        return True
    if _extract_grade_query(normalized) and (
        "ders" in normalized
        or "not" in normalized
        or any(marker in normalized for marker in ("kac", "tane", "aldim", "almis", "almisim", "aldigi", "aldigim"))
    ):
        return True
    if any(marker in normalized for marker in _PERSONAL_TRANSCRIPT_MARKERS):
        return True
    if ("akts" in tokens or "kredi" in tokens) and tokens & {"im", "m", "benim"}:
        return True
    if "ders" in normalized and any(marker in normalized for marker in ("hangi", "neler")):
        if any(marker in normalized for marker in ("aldigim", "aldigi", "aldiklarim", "aldiklari", "almis", "almisim", "aldim")):
            return True
    has_personal_marker = any(
        marker in normalized
        for marker in (
            "benim",
            "bende",
            "aldigim",
            "aldigi",
            "aldiklarim",
            "aldiklari",
            "almis",
            "almisim",
            "aldim",
            "kaldim",
            "gectim",
        )
    )
    has_metric_marker = any(marker in normalized for marker in ("akts", "kredi", "ders", "ortalama"))
    if has_metric_marker and any(marker in normalized for marker in _INSTITUTIONAL_AKTS_MARKERS) and not has_personal_marker:
        return False
    if any(marker in normalized for marker in _TRANSCRIPT_METRIC_MARKERS) and has_personal_marker:
        return True
    return has_personal_marker and has_metric_marker


def is_uploaded_document_followup_query(query: str | None) -> bool:
    normalized = normalize_text(query or "")
    if not normalized:
        return False
    if _forces_uploaded_document_context(normalized):
        return True
    return is_transcript_followup_query(query)


def _forces_uploaded_document_context(normalized: str) -> bool:
    return any(
        marker in normalized
        for marker in (
            "belgeye gore",
            "bu belge",
            "yukledigim belge",
            "yuklenen belge",
            "dosyaya gore",
            "transkripte gore",
            "transkriptime gore",
            "inceleyip",
            "yorumla",
        )
    )


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1254", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def _extract_pdf_text(content: bytes) -> str:
    try:
        import pdfplumber
    except ImportError as exc:
        raise TranscriptProcessingError("PDF metni okumak için pdfplumber kurulu değil.") from exc
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    except Exception as exc:
        raise TranscriptProcessingError("PDF dosyası okunamadı.") from exc


def _extract_pdf_form_fields(content: bytes) -> list[DocumentField]:
    try:
        from PyPDF2 import PdfReader
    except Exception:
        return []
    try:
        reader = PdfReader(io.BytesIO(content))
        raw_fields = reader.get_fields() or {}
    except Exception:
        return []

    fields: list[DocumentField] = []
    for raw_name, raw_field in raw_fields.items():
        if not hasattr(raw_field, "get"):
            continue
        label = _pdf_object_to_text(raw_field.get("/TU"))
        if not label:
            raw_label = _pdf_object_to_text(raw_name)
            if _looks_like_pdf_internal_field_name(raw_label):
                continue
            label = raw_label
        value = _pdf_object_to_text(raw_field.get("/V"))
        if not label:
            continue
        field = _build_document_field(key=label, label=label, value=value or "")
        if field is None:
            continue
        fields.append(
            DocumentField(
                key=field.key,
                label=field.label,
                value=field.value,
                confidence="high" if field.value.strip() else "low",
                state="filled" if field.value.strip() else "empty",
                source="pdf_form",
            )
        )
    return fields


def _looks_like_pdf_internal_field_name(label: str) -> bool:
    normalized = normalize_text(label)
    return bool(re.fullmatch(r"(?:text|fill|toggle|dropdown|sdw)[a-z0-9_]*", normalized))


def _pdf_object_to_text(value: object) -> str:
    if value is None:
        return ""
    text = str(value)
    if text.startswith("/"):
        text = text[1:]
    return text.strip()


def _extract_pdf_ocr_text(content: bytes) -> str:
    try:
        from pdf2image import convert_from_bytes
    except Exception:
        return ""
    try:
        pages = convert_from_bytes(
            content,
            dpi=_OCR_PDF_DPI,
            first_page=1,
            last_page=_OCR_PDF_MAX_PAGES,
            thread_count=1,
        )
    except Exception:
        return ""

    extracted: list[str] = []
    for page in pages[:_OCR_PDF_MAX_PAGES]:
        text = _extract_pil_image_ocr_text(page)
        if text.strip():
            extracted.append(text)
    return "\n".join(extracted)


def _extract_schedule_slots_from_pdf_content(*, filename: str, content: bytes) -> list[dict[str, Any]]:
    try:
        from src.db.schedule_ingest import parse_schedule_slots_from_pdf
    except Exception:
        return []
    suffix = Path(filename or "schedule.pdf").suffix or ".pdf"
    temp_path: str | None = None
    try:
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(content)
            temp_path = tmp.name
        slots = parse_schedule_slots_from_pdf(
            Path(temp_path),
            department=_infer_uploaded_schedule_department(filename),
            term=_infer_uploaded_schedule_term(filename),
            academic_year=_infer_uploaded_schedule_year(filename),
        )
        if slots:
            return slots
        ocr_text = _extract_pdf_ocr_text(content)
        return _extract_schedule_slots_from_text(filename=filename, text=ocr_text)
    except Exception:
        return []
    finally:
        if temp_path:
            try:
                Path(temp_path).unlink(missing_ok=True)
            except Exception:
                pass


def _extract_schedule_slots_from_text(*, filename: str, text: str) -> list[dict[str, Any]]:
    if not _looks_like_course_schedule_document(
        normalized_name=normalize_text(filename or ""),
        normalized_text=normalize_text(text[:4000]),
    ):
        return []

    academic_year = _infer_uploaded_schedule_year(filename) or _infer_uploaded_schedule_year(text) or "bilinmiyor"
    term = _infer_uploaded_schedule_term(filename) or _infer_uploaded_schedule_term(text) or "bilinmiyor"
    department = _infer_uploaded_schedule_department(filename) or _infer_uploaded_schedule_department(text) or "bilinmiyor"
    source_document = filename or "uploaded_schedule"
    slots: list[dict[str, Any]] = []
    current_day: str | None = None
    current_group: str | None = None
    for raw_line in text.splitlines()[:600]:
        line = " ".join(raw_line.split()).strip()
        if len(line) < 8:
            continue
        normalized = normalize_text(line)
        day = _requested_schedule_day(normalized)
        if day:
            current_day = day
        group = _requested_schedule_group(normalized)
        if group:
            current_group = f"{group}. SINIF"
        time_match = re.search(r"\b(\d{1,2}[:.]\d{2})\s*[-–]\s*(\d{1,2}[:.]\d{2})\b", line)
        if not time_match or current_day is None:
            continue
        start = _parse_schedule_time_text(time_match.group(1))
        end = _parse_schedule_time_text(time_match.group(2))
        if start is None or end is None:
            continue
        course_text = _clean_ocr_schedule_course_text(line=line, day=day, group=group, time_text=time_match.group(0))
        if not _looks_like_ocr_schedule_course(course_text):
            continue
        course_code_match = _COURSE_CODE_RE.search(course_text)
        course_code = course_code_match.group(1).replace(" ", "").upper() if course_code_match else None
        course_name = course_text
        classroom = _extract_schedule_classroom_hint(course_text)
        if classroom:
            course_name = course_name.replace(classroom, " ").strip(" -|")
        if course_code:
            course_name = course_name.replace(course_code_match.group(0), " ").strip(" -|") if course_code_match else course_name
        course_name = " ".join(course_name.split())
        if not course_name and course_code:
            course_name = course_code
        if not course_name:
            continue
        slots.append(
            {
                "academic_year": academic_year,
                "term": term,
                "department": department,
                "course_code": course_code,
                "course_key": course_code or normalize_text(course_name),
                "course_name": course_name,
                "schedule_group": current_group,
                "section": None,
                "day_of_week": _display_schedule_day(current_day),
                "start_time": start,
                "end_time": end,
                "classroom": classroom,
                "instructor": None,
                "source_document": source_document,
                "source_url": None,
                "is_active": True,
            }
        )
    deduped: dict[tuple[str, str, str, str, str], dict[str, Any]] = {}
    for slot in slots:
        key = (
            str(slot["day_of_week"]),
            _format_schedule_time(slot["start_time"]),
            _format_schedule_time(slot["end_time"]),
            str(slot.get("schedule_group") or ""),
            str(slot.get("course_key") or ""),
        )
        deduped[key] = slot
    return list(deduped.values())


def _parse_schedule_time_text(value: str) -> clock_time | None:
    cleaned = value.replace(".", ":")
    match = re.match(r"^(\d{1,2}):(\d{2})$", cleaned)
    if not match:
        return None
    hour = int(match.group(1))
    minute = int(match.group(2))
    if not (0 <= hour <= 23 and 0 <= minute <= 59):
        return None
    return clock_time(hour, minute)


def _clean_ocr_schedule_course_text(*, line: str, day: str | None, group: str | None, time_text: str) -> str:
    cleaned = line.replace(time_text, " ")
    if day:
        cleaned = re.sub(day, " ", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(_display_schedule_day(day), " ", cleaned, flags=re.IGNORECASE)
    if group:
        cleaned = re.sub(rf"\b{re.escape(group)}\s*\.?\s*(?:sinif|sınıf|SINIF)?\b", " ", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\b(?:pazartesi|sali|salı|carsamba|çarşamba|persembe|perşembe|cuma|cumartesi|pazar)\b", " ", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\b[1-6]\s*\.?\s*(?:sinif|sınıf|SINIF)\b", " ", cleaned, flags=re.IGNORECASE)
    return " ".join(cleaned.strip(" -:;|").split())


def _looks_like_ocr_schedule_course(value: str) -> bool:
    normalized = normalize_text(value)
    if len(normalized) < 3:
        return False
    if any(marker in normalized for marker in ("saat", "gun", "pazartesi", "sali", "carsamba", "persembe", "cuma")):
        return False
    return any(ch.isalpha() for ch in value)


def _extract_schedule_classroom_hint(value: str) -> str | None:
    tokens = value.split()
    if not tokens:
        return None
    candidates = []
    if len(tokens) >= 2:
        candidates.append(" ".join(tokens[-2:]))
    candidates.append(tokens[-1])
    for candidate in candidates:
        if re.fullmatch(r"(?:MF\s*\d{2,3}|Z\d{1,3}|D\d{1,3}|L\d{1,3}|R\d{1,3}|[A-ZÇĞİÖŞÜ]{1,5}\s*\d{1,3}|Seminer(?:\s+Salonu)?)", candidate, flags=re.IGNORECASE):
            return candidate
    return None


def _infer_uploaded_schedule_department(value: str | None) -> str | None:
    normalized = normalize_text(value or "")
    if "bilgisayar" in normalized and "muh" in normalized:
        return "Bilgisayar Muhendisligi"
    if "elektrik" in normalized and "elektronik" in normalized:
        return "Elektrik-Elektronik Muhendisligi"
    if "fizik" in normalized and ("ogretmen" in normalized or "egitim" in normalized):
        return "Fizik Ogretmenligi"
    return None


def _infer_uploaded_schedule_term(value: str | None) -> str | None:
    normalized = normalize_text(value or "")
    if "bahar" in normalized:
        return "bahar"
    if "guz" in normalized:
        return "guz"
    return None


def _infer_uploaded_schedule_year(value: str | None) -> str | None:
    normalized = normalize_text(value or "")
    match = re.search(r"\b(20\d{2})\s*[-_/]\s*(\d{2,4})\b", normalized)
    if not match:
        return None
    second = match.group(2)
    if len(second) == 2:
        second = match.group(1)[:2] + second
    return f"{match.group(1)}-{second}"


def _extract_docx_text(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise TranscriptProcessingError("DOCX metni okumak için python-docx kurulu değil.") from exc
    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        raise TranscriptProcessingError("DOCX dosyası okunamadı.") from exc
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def _extract_image_ocr_text(content: bytes) -> str:
    try:
        from PIL import Image
    except Exception:
        return ""
    try:
        image = Image.open(io.BytesIO(content))
        return _extract_pil_image_ocr_text(image)
    except Exception:
        return ""


def _extract_pil_image_ocr_text(image) -> str:
    try:
        import pytesseract
    except Exception:
        return ""
    try:
        return str(pytesseract.image_to_string(image, lang="tur+eng") or "")
    except Exception:
        try:
            return str(pytesseract.image_to_string(image, lang="eng") or "")
        except Exception:
            return ""


def _parse_courses(text: str) -> list[TranscriptCourse]:
    courses: list[TranscriptCourse] = []
    seen: set[tuple[str, str, str | None, str]] = set()
    current_term: str | None = None
    for raw_line in text.splitlines():
        line = " ".join(raw_line.split())
        if len(line) < 8:
            continue
        term = _detect_transcript_term_header(line)
        if term is not None:
            current_term = term
        code_match = _COURSE_CODE_RE.search(line)
        if not code_match:
            continue
        course = _parse_course_line(line, code_match)
        if course is None:
            continue
        if current_term and course.term is None:
            course = replace(course, term=current_term)
        key = (normalize_text(course.code), normalize_text(course.name), course.grade, normalize_text(course.term or ""))
        if key in seen:
            continue
        seen.add(key)
        courses.append(course)
    return courses


def _parse_course_line(line: str, code_match: re.Match[str]) -> TranscriptCourse | None:
    code = code_match.group(1).replace(" ", "").upper()
    tail = line[code_match.end() :].strip(" -|")
    if not tail:
        return None
    grade_match = _select_grade_match(list(_GRADE_RE.finditer(tail)))
    grade = grade_match.group(1).upper() if grade_match else _grade_from_status_words(tail)
    number_region = tail[: grade_match.start()] if grade_match else tail
    numbers = [_to_float(match.group(1)) for match in _NUMBER_RE.finditer(number_region)]
    numbers = [number for number in numbers if number is not None]
    credit = numbers[-2] if len(numbers) >= 2 else None
    akts = numbers[-1] if numbers else None
    if credit is not None and not (0 < credit <= _MAX_REASONABLE_COURSE_CREDIT):
        credit = None
    if akts is not None and not (0 < akts <= _MAX_REASONABLE_COURSE_AKTS):
        akts = None
    if credit is None and akts is None and grade is None:
        return None
    name_region = tail[: grade_match.start()] if grade_match else tail
    first_number = _NUMBER_RE.search(name_region)
    if first_number:
        name_region = name_region[: first_number.start()]
    name = name_region.strip(" -|:;")
    if not name or len(name) < 2:
        name = tail[:80].strip(" -|:;")
    status = _status_from_grade_and_text(grade=grade, text=tail)
    return TranscriptCourse(code=code, name=name[:120], credit=credit, akts=akts, grade=grade, status=status)


def _detect_transcript_term_header(line: str) -> str | None:
    normalized = normalize_text(line)
    if not any(marker in normalized for marker in ("donem", "yariyil", "yari yil", "guz", "bahar", "yaz")):
        return None
    if _COURSE_CODE_RE.search(line) and not re.search(r"\b20\d{2}\s*[-/]\s*20\d{2}\b", line):
        return None
    year_match = re.search(r"\b(20\d{2})\s*[-/]\s*(20\d{2})\b", line)
    single_year_match = re.search(r"\b(20\d{2})\b", line)
    season = _extract_season_from_normalized(normalized)
    if season is None and year_match is None:
        return None
    parts: list[str] = []
    if year_match:
        parts.append(f"{year_match.group(1)}-{year_match.group(2)}")
    elif single_year_match:
        parts.append(single_year_match.group(1))
    if season:
        parts.append(season)
    return " ".join(parts) if parts else None


def _extract_season_from_normalized(normalized: str) -> str | None:
    if "bahar" in normalized:
        return "Bahar"
    if "guz" in normalized:
        return "Güz"
    if re.search(r"\byaz\b", normalized):
        return "Yaz"
    return None


def _select_grade_match(matches: list[re.Match[str]]) -> re.Match[str] | None:
    detailed_matches = [match for match in matches if match.group(1).upper() in _DETAILED_GRADES]
    if detailed_matches:
        return detailed_matches[-1]
    return matches[-1] if matches else None


def _to_float(value: str | None) -> float | None:
    if value is None:
        return None
    try:
        return float(value.replace(",", "."))
    except ValueError:
        return None


def _grade_from_status_words(text: str) -> str | None:
    normalized = normalize_text(text)
    if "basarisiz" in normalized or "kaldi" in normalized:
        return "FF"
    if "basarili" in normalized or "gecti" in normalized:
        return "G"
    return None


def _status_from_grade_and_text(*, grade: str | None, text: str) -> str:
    normalized = normalize_text(text)
    if "basarisiz" in normalized or "kaldi" in normalized:
        return "failed"
    if "basarili" in normalized or "gecti" in normalized:
        return "passed"
    if not grade:
        return "unknown"
    if grade.upper() in _FAILED_GRADES:
        return "failed"
    if grade.upper() in _PASSED_GRADES:
        return "passed"
    return "unknown"


def _extract_total_value(text: str, *, kind: str) -> float | None:
    label_patterns = (
        (r"(?:toplam|kazanılan|kazanilan|tamamlanan|başarılan|basarilan)\s+(?:akts|ects)", "akts"),
        (r"(?:akts|ects)\s+(?:toplamı|toplami|toplam|kazanılan|kazanilan)", "akts"),
        (r"(?:toplam|kazanılan|kazanilan|tamamlanan)\s+(?:kredi|ulusal kredi)", "credit"),
        (r"(?:kredi|ulusal kredi)\s+(?:toplamı|toplami|toplam|kazanılan|kazanilan)", "credit"),
    )
    candidates: list[float] = []
    for label, label_kind in label_patterns:
        if label_kind != kind:
            continue
        pattern = re.compile(label + r"[^0-9]{0,30}(\d{1,3}(?:[,.]\d{1,2})?)(?!\d)", re.IGNORECASE)
        candidates.extend(
            value
            for value in (_to_float(match.group(1)) for match in pattern.finditer(text))
            if value is not None
        )
    if not candidates:
        return None
    return max(candidates)


def _extract_gpa(text: str) -> float | None:
    patterns = (
        r"\b(?:gno|gano|agno|genel not ortalaması|genel not ortalamasi)\b[^0-9]{0,30}(\d{1,2}[,.]\d{1,2})",
        r"\b(?:cgpa|gpa)\b[^0-9]{0,30}(\d{1,2}[,.]\d{1,2})",
    )
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            return _to_float(match.group(1))
    return None


def _classify_document(*, filename: str, text: str) -> str:
    normalized_name = normalize_text(filename or "")
    normalized_text = normalize_text(text[:4000])
    if _looks_like_course_schedule_document(normalized_name=normalized_name, normalized_text=normalized_text):
        return "schedule_document"
    if "adli sicil" in normalized_name or "sicil kaydi" in normalized_name or "adli sicil" in normalized_text:
        return "criminal_record"
    if any(marker in normalized_text for marker in ("ogrenci belgesi", "ogrencilik belgesi", "ogrencilik durumu")):
        return "student_document"
    if any(marker in normalized_name for marker in ("ogrenci belgesi", "ogrencilik belgesi")):
        return "student_document"
    if any(marker in normalized_text for marker in ("sonuc belgesi", "sinav sonucu", "sorgulama sonucu")):
        return "certificate_or_result"
    if any(marker in normalized_text for marker in ("yonerge", "talimat", "kilavuz", "uygulama esas")):
        return "policy_or_instruction"
    if any(marker in normalized_name for marker in ("form", "dilekce", "basvuru")):
        return "form_document"
    if any(marker in normalized_text for marker in ("staj basvuru formu", "is basvuru formu", "aday bilgi formu", "basvuru formu")):
        return "form_document"
    if any(marker in normalized_name for marker in ("transkript", "not dokum", "not dokumu")):
        return "transcript"
    if any(marker in normalized_text for marker in ("transkript", "not dokum", "not dokumu", "gno", "agno")):
        return "transcript"
    course_count = len(_parse_courses(text))
    has_transcript_metric = any(marker in normalized_text for marker in ("toplam akts", "toplam kredi", "genel not ortalamasi"))
    detailed_grade_hits = len(re.findall(r"\b(?:AA|BA|BB|CB|CC|DC|DD|FD|FF|YT|YZ)\b", text, flags=re.IGNORECASE))
    if course_count >= 2 and (has_transcript_metric or detailed_grade_hits >= 3):
        return "transcript"
    if len(normalized_text.strip()) < 12:
        return "unknown_document"
    return "generic_document"


def _looks_like_course_schedule_document(*, normalized_name: str, normalized_text: str) -> bool:
    haystack = f"{normalized_name}\n{normalized_text}".replace("_", " ").replace("-", " ")
    schedule_markers = (
        "ders program",
        "ders programi",
        "lisans ders program",
        "haftalik ders",
        "haftalik program",
        "bahar ders program",
        "guz ders program",
    )
    if any(marker in haystack for marker in schedule_markers):
        return True
    weekday_hits = sum(
        1
        for marker in ("pazartesi", "sali", "carsamba", "persembe", "cuma")
        if marker in normalized_text
    )
    time_slot_count = len(re.findall(r"\b\d{1,2}[:.]\d{2}\s*[-–]\s*\d{1,2}[:.]\d{2}\b", normalized_text))
    room_markers = ("derslik", "mf", "lab", "sinif", "sınıf")
    if weekday_hits >= 2 and (time_slot_count >= 2 or any(marker in normalized_text for marker in room_markers)):
        return True
    return False


def _build_document_facts(
    *,
    filename: str,
    text: str,
    document_type: str,
    fields: tuple[DocumentField, ...],
    schedule_slots: tuple[dict[str, Any], ...],
    courses: tuple[TranscriptCourse, ...],
    total_akts: float | None,
    total_credit: float | None,
    gpa: float | None,
    program_name: str | None,
    extraction_confidence: str,
    warnings: tuple[str, ...],
) -> DocumentFacts:
    synthetic_fields: list[DocumentField] = []
    if program_name:
        synthetic_fields.append(
            DocumentField(key="program", label="Program", value=program_name, confidence="medium", source="transcript_parser")
        )
    if total_akts is not None:
        synthetic_fields.append(
            DocumentField(
                key="toplam akts",
                label="Toplam AKTS",
                value=_format_number(total_akts),
                confidence="medium",
                source="transcript_parser",
            )
        )
    if total_credit is not None:
        synthetic_fields.append(
            DocumentField(
                key="toplam kredi",
                label="Toplam kredi",
                value=_format_number(total_credit),
                confidence="medium",
                source="transcript_parser",
            )
        )
    if gpa is not None:
        synthetic_fields.append(
            DocumentField(key="gno agno", label="GNO/AGNO", value=_format_number(gpa), confidence="medium", source="transcript_parser")
        )

    all_fields = tuple(fields) + tuple(synthetic_fields)
    tables: list[dict[str, Any]] = []
    if schedule_slots:
        tables.append({"kind": "schedule_slots", "row_count": len(schedule_slots)})
    if courses:
        tables.append({"kind": "transcript_courses", "row_count": len(courses)})
    person_name = _extract_person_name_from_fields(all_fields)
    return DocumentFacts(
        document_type=document_type,
        title=_extract_document_title(filename=filename, text=text, document_type=document_type),
        issuer=_extract_document_issuer(text),
        person_name=person_name,
        fields=all_fields,
        tables=tuple(tables),
        schedule_slots=schedule_slots,
        transcript_courses=courses,
        dates=tuple(field for field in all_fields if "tarih" in field.key or "date" in field.key),
        status=_extract_document_status(document_type=document_type, text=text),
        extraction_confidence=extraction_confidence,
        warnings=warnings,
        raw_text_preview=" ".join(text.split())[:1200],
    )


def _document_facts(document: TranscriptDocument) -> DocumentFacts:
    if document.facts is not None:
        return document.facts
    return _build_document_facts(
        filename=document.filename,
        text=document.text,
        document_type=document.document_type,
        fields=document.fields,
        schedule_slots=document.schedule_slots,
        courses=document.courses,
        total_akts=document.total_akts,
        total_credit=document.total_credit,
        gpa=document.gpa,
        program_name=document.program_name,
        extraction_confidence=document.parse_confidence,
        warnings=document.warnings,
    )


def _extract_document_title(*, filename: str, text: str, document_type: str) -> str | None:
    for raw_line in text.splitlines()[:12]:
        line = " ".join(raw_line.split()).strip(" -:;|")
        if len(line) < 4 or len(line) > 140:
            continue
        normalized = normalize_text(line)
        if any(marker in normalized for marker in ("http", "www", "sayfa")):
            continue
        if document_type == "schedule_document" and "ders program" in normalized:
            return line
        if document_type == "transcript" and any(marker in normalized for marker in ("transkript", "not dokum")):
            return line
        if document_type != "generic_document":
            return line
    stem = Path(filename or "").stem.replace("_", " ").replace("-", " ").strip()
    return stem or None


def _extract_document_issuer(text: str) -> str | None:
    normalized = normalize_text(text[:3000])
    if "ondokuz mayis universitesi" in normalized or "on dokuz mayis universitesi" in normalized:
        return "Ondokuz Mayis Universitesi"
    if "adalet bakanligi" in normalized:
        return "T.C. Adalet Bakanligi"
    if "yuksekogretim kurulu" in normalized:
        return "Yuksekogretim Kurulu"
    if "havelsan" in normalized:
        return "HAVELSAN"
    return None


def _extract_person_name_from_fields(fields: tuple[DocumentField, ...]) -> str | None:
    for key in ("adi soyadi", "ad soyad", "ad soyadi"):
        for field in fields:
            if _canonical_document_field_key(field.key) == key and field.state == "filled" and field.value.strip():
                return field.value.strip()
    return None


def _extract_document_status(*, document_type: str, text: str) -> str | None:
    normalized = normalize_text(text)
    if document_type == "criminal_record":
        no_record = "adli sicil kaydi yoktur" in normalized
        no_archive = "adli sicil arsiv kaydi yoktur" in normalized
        if no_record and no_archive:
            return "adli_sicil_ve_arsiv_kaydi_yok"
        if no_record:
            return "adli_sicil_kaydi_yok"
        if "adli sicil kaydi vardir" in normalized:
            return "adli_sicil_kaydi_var"
    if document_type == "student_document":
        if "aktif ogrenci" in normalized or "ogrencilik durumu aktif" in normalized:
            return "aktif_ogrenci"
        if "pasif" in normalized:
            return "pasif_ogrenci"
    return None


def _facts_payload(document: TranscriptDocument) -> dict[str, Any]:
    facts = _document_facts(document)
    return {
        "document_type": facts.document_type,
        "title": facts.title,
        "issuer": facts.issuer,
        "person_name": facts.person_name,
        "status": facts.status,
        "extraction_confidence": facts.extraction_confidence,
        "warnings": list(facts.warnings),
        "fields": [
            {
                "label": field.label,
                "key": field.key,
                "value": field.value,
                "state": field.state,
                "source": field.source,
                "confidence": field.confidence,
            }
            for field in facts.fields[:60]
        ],
        "tables": list(facts.tables),
        "schedule_slot_count": len(facts.schedule_slots),
        "transcript_course_count": len(facts.transcript_courses),
        "dates": [{"label": field.label, "value": field.value, "state": field.state} for field in facts.dates[:20]],
        "raw_text_preview": facts.raw_text_preview,
    }


def _build_extraction_debug(
    *,
    text: str,
    document_type: str,
    extraction_mode: str,
    warnings: list[str],
    fields: tuple[DocumentField, ...],
    courses: tuple[TranscriptCourse, ...],
    parse_confidence: str,
) -> dict[str, Any]:
    filled_fields = [field for field in fields if field.state == "filled" and field.value.strip()]
    empty_fields = [field for field in fields if field.state != "filled" or not field.value.strip()]
    return {
        "document_type": document_type,
        "extraction_mode": extraction_mode,
        "parse_confidence": parse_confidence,
        "raw_text_chars": len(text),
        "line_count": len(text.splitlines()),
        "field_count": len(fields),
        "filled_field_count": len(filled_fields),
        "empty_field_count": len(empty_fields),
        "course_count": len(courses),
        "facts_enabled": True,
        "table_like_label_line_count": _count_table_like_label_lines(text),
        "warnings": list(warnings),
    }


def _count_table_like_label_lines(text: str) -> int:
    count = 0
    for raw_line in text.splitlines()[:260]:
        line = " ".join(raw_line.split()).strip()
        if len(_known_labels_in_line(line)) >= 2:
            count += 1
    return count


def _extract_document_fields(text: str) -> list[DocumentField]:
    fields: list[DocumentField] = []
    seen: dict[str, int] = {}
    for raw_line in text.splitlines()[:260]:
        line = " ".join(raw_line.split()).strip()
        if not _looks_like_field_line(line):
            continue
        key: str | None = None
        label: str | None = None
        value: str | None = None

        separator_match = _FIELD_SEPARATOR_RE.match(line)
        if separator_match:
            label = separator_match.group(1).strip(" -:;|")
            value = separator_match.group(2).strip(" -:;|")
            key = _field_key(label)
        else:
            spaced_match = _FIELD_SPACED_RE.match(line)
            if spaced_match:
                label = spaced_match.group(1).strip(" -:;|")
                value = spaced_match.group(2).strip(" -:;|")
                key = _field_key(label)
            else:
                known = _known_label_value_from_line(line)
                if known is not None:
                    key, label, value = known

        field = _build_document_field(key=key, label=label, value=value)
        if field is not None and field.key in seen:
            existing = fields[seen[field.key]]
            if existing.state != "filled" and field.state == "filled":
                fields[seen[field.key]] = field
            field = None
        if field is None:
            for label_key, display_label in _known_labels_in_line(line):
                if label_key in seen:
                    continue
                seen[label_key] = len(fields)
                fields.append(
                    DocumentField(
                        key=label_key,
                        label=display_label,
                        value="",
                        confidence="low",
                        state="empty",
                        source="label_inventory",
                    )
                )
                if len(fields) >= 80:
                    break
            if len(fields) >= 80:
                break
            continue
        seen[field.key] = len(fields)
        fields.append(field)
        for label_key, display_label in _known_labels_in_line(line):
            if label_key in seen:
                continue
            seen[label_key] = len(fields)
            fields.append(
                DocumentField(
                    key=label_key,
                    label=display_label,
                    value="",
                    confidence="low",
                    state="empty",
                    source="label_inventory",
                )
            )
            if len(fields) >= 80:
                break
        if len(fields) >= 80:
            break
    return fields


def _merge_document_fields(*field_groups: tuple[DocumentField, ...]) -> tuple[DocumentField, ...]:
    merged: list[DocumentField] = []
    seen: dict[str, int] = {}
    duplicate_counts: dict[str, int] = {}
    for group in field_groups:
        for field in group:
            key = _canonical_document_field_key(field.key)
            if key not in seen:
                seen[key] = len(merged)
                merged.append(field)
                continue
            existing = merged[seen[key]]
            if existing.state != "filled" and field.state == "filled":
                merged[seen[key]] = field
                continue
            if existing.state == "filled" and field.state == "filled" and existing.value.strip() != field.value.strip():
                duplicate_counts[key] = duplicate_counts.get(key, 1) + 1
                duplicate_key = f"{key} {duplicate_counts[key]}"
                seen[duplicate_key] = len(merged)
                merged.append(
                    DocumentField(
                        key=duplicate_key,
                        label=f"{field.label} {duplicate_counts[key]}",
                        value=field.value,
                        confidence=field.confidence,
                        state=field.state,
                        source=field.source,
                    )
                )
    return tuple(merged)


def _canonical_document_field_key(key: str) -> str:
    normalized = _field_key(key)
    for canonical, aliases in _FIELD_LABEL_ALIASES:
        if normalized == canonical or normalized in aliases:
            return canonical
    return normalized


def _looks_like_field_line(line: str) -> bool:
    if not line or len(line) < 4 or len(line) > 220:
        return False
    normalized = normalize_text(line)
    if any(marker in normalized for marker in _FIELD_KEY_SKIP_MARKERS):
        return False
    if _COURSE_CODE_RE.search(line):
        return False
    return any(ch.isalpha() for ch in line)


def _known_label_value_from_line(line: str) -> tuple[str, str, str] | None:
    normalized = _field_key(line)
    for label_key in _iter_known_label_keys():
        if not (normalized == label_key or normalized.startswith(label_key + " ")):
            continue
        value = line[len(label_key) :].strip(" -:;|")
        normalized_value = _field_key(value)
        if any(
            other != label_key and _label_occurs_in_line(other, normalized_value)
            for other in _iter_known_label_keys()
        ):
            value = ""
        return label_key, _display_label(label_key), value
    return None


def _known_labels_in_line(line: str) -> list[tuple[str, str]]:
    normalized = _field_key(line)
    found: list[tuple[str, str]] = []
    for label_key in _iter_known_label_keys():
        if _label_occurs_in_line(label_key, normalized):
            found.append((label_key, _display_label(label_key)))
    return found


def _iter_known_label_keys() -> tuple[str, ...]:
    keys = list(_KNOWN_FIELD_LABELS)
    keys.extend(canonical for canonical, _aliases in _FIELD_LABEL_ALIASES)
    return tuple(dict.fromkeys(keys))


def _label_occurs_in_line(label_key: str, normalized_line: str) -> bool:
    if not label_key or label_key in normalized_line:
        return bool(label_key)
    for canonical, aliases in _FIELD_LABEL_ALIASES:
        if label_key == canonical and any(alias in normalized_line for alias in aliases):
            return True
    return False


def _build_document_field(*, key: str | None, label: str | None, value: str | None) -> DocumentField | None:
    if not key or not label or value is None:
        return None
    key = _field_key(key)
    label = label.strip(" -:;|")
    value = value.strip(" -:;|")
    if not key or len(key) < 2:
        return None
    if _looks_like_known_field_label(value, current_key=key):
        value = ""
    if value and len(_field_key(value)) < 3 and not any(ch.isdigit() for ch in value):
        value = ""
    if len(label) > 60 or len(value) > 180:
        return None
    normalized_label = normalize_text(label)
    normalized_value = normalize_text(value)
    if any(marker in normalized_label for marker in _FIELD_KEY_SKIP_MARKERS):
        return None
    if value and normalized_label == normalized_value:
        return None
    return DocumentField(
        key=key,
        label=label,
        value=value,
        confidence="medium" if value else "low",
        state="filled" if value else "empty",
        source="text_layer",
    )


def _looks_like_known_field_label(value: str, *, current_key: str | None = None) -> bool:
    normalized = _field_key(value)
    if not normalized:
        return False
    return any(
        label_key != current_key and (normalized == label_key or _label_occurs_in_line(label_key, normalized))
        for label_key in _iter_known_label_keys()
    )


def _field_key(label: str) -> str:
    normalized = normalize_text(label)
    return _FIELD_KEY_CLEAN_RE.sub(" ", normalized).strip()


def _display_label(key: str) -> str:
    return key.replace(" no", " No").replace(" tc ", " TC ").title()


def _extract_program_name(text: str) -> str | None:
    for raw_line in text.splitlines():
        line = " ".join(raw_line.split()).strip()
        if not line:
            continue
        normalized = normalize_text(line)
        if not any(marker in normalized for marker in ("program", "bolum", "ana bilim", "anabilim")):
            continue
        candidate = _candidate_after_program_label(line)
        if _is_plausible_program_name(candidate):
            return candidate[:90]

    suffix_pattern = re.compile(
        r"\b([A-ZÇĞİÖŞÜ][A-Za-zÇĞİÖŞÜçğıöşü\s\-]{2,80}"
        r"(?:Mühendisliği|Muhendisligi|Öğretmenliği|Ogretmenligi|Eğitimi|Egitimi))\b",
        flags=re.IGNORECASE,
    )
    for match in suffix_pattern.finditer(text):
        candidate = " ".join(match.group(1).split()).strip(" -:;|")
        if _is_plausible_program_name(candidate):
            return candidate[:90]
    return None


def _candidate_after_program_label(line: str) -> str:
    if ":" in line:
        return line.split(":", 1)[1].strip(" -:;|")
    tokens = line.split()
    for index, token in enumerate(tokens):
        normalized_token = normalize_text(token)
        if normalized_token.startswith(("program", "bolum", "anabilim")):
            return " ".join(tokens[index + 1 :]).strip(" -:;|()")
        if normalized_token == "ana" and index + 1 < len(tokens) and normalize_text(tokens[index + 1]).startswith("bilim"):
            return " ".join(tokens[index + 2 :]).strip(" -:;|()")
    return line.strip(" -:;|")


def _is_plausible_program_name(candidate: str) -> bool:
    normalized = normalize_text(candidate)
    if len(candidate.strip()) < 3:
        return False
    forbidden = (
        "kredi turu",
        "akts",
        "anadal",
        "program turu",
        "kayit tarihi",
        "giris tarihi",
        "tarih",
        "gno",
        "agno",
        "ogrenci no",
        "tc kimlik",
    )
    if any(marker in normalized for marker in forbidden):
        return False
    if re.search(r"\d{4}", candidate):
        return False
    return any(
        marker in normalized
        for marker in (
            "muhendisligi",
            "ogretmenligi",
            "egitimi",
            "fakultesi",
            "myo",
            "onlisans",
            "lisans",
            "tip",
            "dis hekimligi",
            "eczacilik",
            "veteriner",
        )
    )

def _is_reasonable_total(value: float | None, *, kind: str) -> bool:
    if value is None:
        return True
    if value <= 0:
        return False
    limit = _MAX_REASONABLE_TRANSCRIPT_AKTS if kind == "akts" else _MAX_REASONABLE_TRANSCRIPT_CREDIT
    return value <= limit


def _has_unreasonable_total_literal(text: str, *, kind: str) -> bool:
    label = r"(?:akts|ects)" if kind == "akts" else r"(?:kredi|ulusal kredi)"
    pattern = re.compile(
        r"(?:toplam|kazanÄ±lan|kazanilan|tamamlanan|baÅŸarÄ±lan|basarilan)\s+"
        + label
        + r"[^0-9]{0,30}(\d{4,})",
        re.IGNORECASE,
    )
    return bool(pattern.search(text))


def _sum_reliable_course_metric(courses: tuple[TranscriptCourse, ...], *, metric: str) -> float | None:
    passed = [course for course in courses if course.status == "passed"]
    candidates = passed if passed else [course for course in courses if course.status != "failed"]
    total = 0.0
    for course in candidates:
        value = course.akts if metric == "akts" else course.credit
        if value is None:
            continue
        if metric == "akts" and not (0 < value <= _MAX_REASONABLE_COURSE_AKTS):
            continue
        if metric == "credit" and not (0 < value <= _MAX_REASONABLE_COURSE_CREDIT):
            continue
        total += value
    if not _is_reasonable_total(total, kind=metric):
        return None
    return total if total > 0 else None


def _transcript_parse_confidence(
    *,
    courses: tuple[TranscriptCourse, ...],
    total_akts: float | None,
    total_credit: float | None,
) -> str:
    if len(courses) >= 8 and (total_akts is not None or total_credit is not None):
        return "high"
    if len(courses) >= 2 or total_akts is not None or total_credit is not None:
        return "medium"
    return "low"


def _answer_transcript_question(*, document: TranscriptDocument, query: str) -> str | None:
    normalized = normalize_text(query)
    if _asks_graduation_progress(normalized):
        return _answer_transcript_graduation_progress(document=document)
    if any(marker in normalized for marker in ("not dagilimi", "harf notu dagilimi", "notlarim nasil")):
        return _format_grade_distribution(document)
    grade_query = _extract_grade_query(normalized)
    if grade_query and (
        "ders" in normalized
        or "not" in normalized
        or any(marker in normalized for marker in ("kac", "tane", "aldim", "almisim", "aldigim", "var", "suru", "cok", "emin"))
    ):
        return _answer_grade_query(document=document, grade=grade_query, normalized_query=normalized)
    if any(marker in normalized for marker in ("harc", "borc", "ucret", "odeme")):
        return "Bu transkriptte harç/borç/ödeme bilgisi tespit edemedim."
    if any(
        marker in normalized
        for marker in ("hangi ders", "aldigim ders", "aldigi ders", "aldiklarim", "aldiklari", "almis", "almisim", "aldim")
    ):
        term_request = _extract_transcript_term_request(normalized)
        if term_request is not None:
            return _answer_term_course_list(document=document, term_request=term_request)
        return _format_course_list(
            title="Transkriptte görünen aldığınız dersler",
            courses=document.courses,
            empty="Transkriptte ders satırı tespit edemedim.",
        )
    if any(marker in normalized for marker in ("kaldi", "kaldigim", "basarisiz", "ff", "fd")):
        return _format_course_list(
            title="Transkriptte başarısız görünen dersler",
            courses=document.failed_courses,
            empty="Transkriptte başarısız ders tespit edemedim.",
        )
    if any(marker in normalized for marker in ("gectigim", "basarili", "gecen ders", "geçtiğim")):
        return _format_course_list(
            title="Transkriptte başarılı görünen dersler",
            courses=document.passed_courses,
            empty="Transkriptte başarılı ders tespit edemedim.",
        )
    if "akts" in normalized:
        if document.total_akts is None:
            return "Transkriptte toplam AKTS bilgisini güvenilir biçimde çıkaramadım."
        return f"Transkriptte görünen toplam AKTS: {_format_number(document.total_akts)}."
    if "kredi" in normalized:
        if document.total_credit is None:
            return "Transkriptte toplam kredi bilgisini güvenilir biçimde çıkaramadım."
        return f"Transkriptte görünen toplam kredi: {_format_number(document.total_credit)}."
    if any(marker in normalized for marker in ("ortalama", "gno", "agno", "gano")):
        historical_term_answer = _answer_historical_term_gpa_question(document=document, normalized_query=normalized)
        if historical_term_answer is not None:
            return historical_term_answer
        interpretive = _answer_gpa_interpretation(document=document, normalized_query=normalized)
        if interpretive is not None:
            return interpretive
        if document.gpa is None:
            return "Transkriptte genel not ortalaması bilgisini güvenilir biçimde çıkaramadım."
        return f"Transkriptte görünen genel not ortalaması: {_format_number(document.gpa)}."
    if any(marker in normalized for marker in ("ozet", "durum", "ne gorunuyor", "neler var")):
        return _format_transcript_summary(document)
    return None


def _extract_transcript_term_request(normalized_query: str) -> _TermRequest | None:
    season = _extract_season_from_normalized(normalized_query)
    academic_year: str | None = None
    single_year: int | None = None
    year_match = re.search(r"\b(20\d{2})\s*[-/]\s*(20\d{2})\b", normalized_query)
    if year_match:
        first = int(year_match.group(1))
        second = int(year_match.group(2))
        if first != second:
            academic_year = f"{first}-{second}"
        else:
            single_year = first
    elif (single_year_match := re.search(r"\b(20\d{2})\b", normalized_query)):
        single_year = int(single_year_match.group(1))

    has_term_marker = any(marker in normalized_query for marker in ("donem", "yariyil", "yari yil", "guz", "bahar", "yaz"))
    if not has_term_marker and academic_year is None and single_year is None:
        return None

    display_parts: list[str] = []
    if academic_year:
        display_parts.append(academic_year)
    elif single_year:
        display_parts.append(str(single_year))
    if season:
        display_parts.append(season)
    display = " ".join(display_parts) if display_parts else "istenen dönem"
    return _TermRequest(
        academic_year=academic_year,
        season=season,
        display=display,
        single_year=single_year,
    )


def _answer_term_course_list(*, document: TranscriptDocument, term_request: _TermRequest) -> str:
    courses_with_term = tuple(course for course in document.courses if course.term)
    if not courses_with_term:
        return (
            "Transkriptte dersler görünüyor; ancak ders satırlarını dönem/yıl bilgisiyle güvenilir biçimde "
            f"eşleştiremedim. Bu yüzden {term_request.display} dönemindeki dersleri ayrı listeleyemiyorum."
        )

    matched = tuple(course for course in courses_with_term if _course_matches_term_request(course, term_request))
    if not matched:
        return f"Transkriptte {term_request.display} dönemi için ders satırı tespit edemedim."

    return _format_course_list(
        title=f"Transkriptte {term_request.display} döneminde görünen aldığınız dersler",
        courses=matched,
        empty=f"Transkriptte {term_request.display} dönemi için ders satırı tespit edemedim.",
    )


def _course_matches_term_request(course: TranscriptCourse, term_request: _TermRequest) -> bool:
    if not course.term:
        return False
    normalized_term = normalize_text(course.term)
    if term_request.season and normalize_text(term_request.season) not in normalized_term:
        return False
    if term_request.academic_year:
        return term_request.academic_year in course.term
    if term_request.single_year is not None:
        return str(term_request.single_year) in course.term
    return True


def _format_grade_distribution(document: TranscriptDocument) -> str:
    counts: dict[str, int] = {}
    for course in document.courses:
        if not course.grade:
            continue
        grade = course.grade.upper()
        counts[grade] = counts.get(grade, 0) + 1
    if not counts:
        return "Transkriptte harf notu dağılımını güvenilir biçimde çıkaramadım."
    order = ("AA", "BA", "BB", "CB", "CC", "DC", "DD", "FD", "FF", "YT", "YZ", "G", "K", "P", "F")
    lines = ["Transkriptte görünen harf notu dağılımı:"]
    for grade in order:
        if grade in counts:
            lines.append(f"- {grade}: {counts[grade]} ders")
    for grade in sorted(set(counts) - set(order)):
        lines.append(f"- {grade}: {counts[grade]} ders")
    return "\n".join(lines)


def _extract_grade_query(normalized: str) -> str | None:
    tokens = set(normalized.split())
    for grade in ("aa", "ba", "bb", "cb", "cc", "dc", "dd", "fd", "ff", "yt", "yz"):
        if grade in tokens:
            return grade.upper()
    for grade in ("g", "k", "p", "f"):
        if grade in tokens and ("not" in normalized or "ders" in normalized):
            return grade.upper()
    return None


def _answer_grade_query(*, document: TranscriptDocument, grade: str, normalized_query: str) -> str:
    courses = tuple(course for course in document.courses if (course.grade or "").upper() == grade)
    if any(marker in normalized_query for marker in ("hangi", "liste", "neler")):
        return _format_course_list(
            title=f"Transkriptte {grade} notlu dersler",
            courses=courses,
            empty=f"Transkriptte {grade} notlu ders tespit edemedim.",
        )
    count = len(courses)
    answer = f"Transkriptte {grade} notlu {count} ders görünüyor."
    if count and count <= 8:
        names = ", ".join(f"{course.code} {course.name}" for course in courses)
        answer += f" Dersler: {names}."
    return answer


def _asks_graduation_progress(normalized: str) -> bool:
    return any(marker in normalized for marker in ("mezun", "mezuniyet", "tamamlamam")) and any(
        marker in normalized for marker in ("akts", "kredi", "gerekiyor", "gerekir", "kac")
    )


def _answer_transcript_graduation_progress(*, document: TranscriptDocument) -> str:
    from src.agents.student.graduation_utils import resolve_program_level

    completed = document.total_akts
    program_text = document.program_name or ""
    level = resolve_program_level(program_text, student_department=program_text).get("education_level")
    required: int | None = None
    level_text = "program"
    if level == "associate":
        required = 120
        level_text = "ön lisans"
    elif level == "bachelor":
        required = 240
        level_text = "normal dört yıllık lisans"
    elif level == "special_long":
        level_text = "özel uzun program"

    intro = "Belgeye göre"
    if document.program_name:
        intro += f" programınız {document.program_name} görünüyor"
    else:
        intro += " program adını güvenilir biçimde çıkaramadım"

    if completed is None:
        return (
            f"{intro}. Transkriptte tamamlanan toplam AKTS bilgisini güvenilir biçimde çıkaramadım; "
            "bu yüzden belgeye göre kalan AKTS hesabı yapamıyorum."
        )
    if required is None:
        return (
            f"{intro}. Toplam AKTS: {_format_number(completed)}. "
            f"Bu {level_text} için toplam mezuniyet AKTS koşulunu bu belgeden doğrulayamadım; resmi program koşuluyla kontrol etmek gerekir."
        )
    remaining = max(required - completed, 0)
    return (
        f"{intro}. Toplam AKTS: {_format_number(completed)}. "
        f"{level_text} programlarda mezuniyet hedefi toplam {required} AKTS olduğu için "
        f"yaklaşık {_format_number(float(remaining))} AKTS daha tamamlamanız gerekir."
    )


def _format_transcript_summary(document: TranscriptDocument) -> str:
    parts = [
        f"Transkript okundu: {document.filename}",
        f"Tespit edilen ders sayısı: {len(document.courses)}",
    ]
    if document.program_name:
        parts.append(f"Program: {document.program_name}")
    if document.total_akts is not None:
        parts.append(f"Toplam AKTS: {_format_number(document.total_akts)}")
    if document.total_credit is not None:
        parts.append(f"Toplam kredi: {_format_number(document.total_credit)}")
    if document.gpa is not None:
        parts.append(f"GNO/AGNO: {_format_number(document.gpa)}")
    failed_count = len(document.failed_courses)
    if failed_count:
        parts.append(f"Başarısız görünen ders: {failed_count}")
    if document.warnings:
        parts.append("Not: " + " ".join(document.warnings))
    return "\n".join(f"- {part}" for part in parts)


def _format_generic_document_summary(document: TranscriptDocument) -> str:
    facts = _document_facts(document)
    preview = " ".join(document.text.split())[:500]
    if not preview:
        return "Belgeden okunabilir metin ??kar?ld? ancak ?zetlenebilecek i?erik bulunamad?."
    lines = [
        f"Belge okundu: {document.filename}",
        f"- Tur: {_display_document_type(facts.document_type)}",
        f"- Metin ??karma y?ntemi: {document.extraction_mode}",
    ]
    if facts.title:
        lines.append(f"- Baslik: {facts.title}")
    if facts.issuer:
        lines.append(f"- Kurum: {facts.issuer}")
    if facts.person_name:
        lines.append(f"- Kisi: {facts.person_name}")
    if facts.status:
        lines.append(f"- Durum/sonuc: {facts.status}")
    if facts.tables:
        lines.append("- Tablolar: " + ", ".join(f"{table.get('kind')} ({table.get('row_count')})" for table in facts.tables))
    if document.fields:
        field_preview = ", ".join(
            f"{field.label}: {field.value}" if field.value.strip() else f"{field.label}: [bos]"
            for field in document.fields[:8]
        )
        lines.append(f"- Tespit edilen alanlar: {field_preview}")
    lines.append(f"- K?sa i?erik: {preview}")
    return "\n".join(lines)


def _display_document_type(document_type: str) -> str:
    return {
        "transcript": "transkript",
        "schedule_document": "ders programi belgesi",
        "form_document": "form/basvuru belgesi",
        "student_document": "ogrenci belgesi",
        "criminal_record": "adli sicil/sonuc belgesi",
        "certificate_or_result": "sonuc belgesi",
        "policy_or_instruction": "yonerge/talimat belgesi",
        "unknown_document": "bilinmeyen belge",
        "generic_document": "genel belge",
    }.get(document_type, "genel belge")


def _answer_schedule_document_question(
    *,
    document: TranscriptDocument,
    query: str,
    document_intent: DocumentQuestionIntent | None,
) -> str | None:
    if document.document_type != "schedule_document":
        return None
    normalized = normalize_text(query)
    if document_intent == DocumentQuestionIntent.DOCUMENT_SUMMARY or _is_open_generic_document_question(normalized):
        return _format_schedule_document_summary(document)
    if document_intent == DocumentQuestionIntent.HYBRID_DOCUMENT_RAG or _looks_like_schedule_detail_query(normalized):
        if not document.schedule_slots:
            return (
                "Yüklenen ders programı belgesi görsel/tablo formatında görünüyor; bu belgeden gün-saat-sınıf satırlarını "
                "güvenilir biçimde ayrıştıramadım. Sistemde yapılandırılmış ders programı kaydı varsa cevap için onu kullanmak gerekir; "
                "kayıt yoksa bu belge üzerinden net ders/saat bilgisi uyduramam."
            )
        return _answer_schedule_slots_question(slots=document.schedule_slots, query=query)
    return None


def _answer_historical_term_gpa_question(*, document: TranscriptDocument, normalized_query: str) -> str | None:
    asks_historical_term = any(
        marker in normalized_query
        for marker in (
            "hangi donem",
            "hangi donemlerde",
            "kac donem",
            "donemleri",
            "donemlerde",
            "bitirmis",
            "bitirmisim",
            "bitirdim",
            "bitirmis miyim",
        )
    )
    if not asks_historical_term:
        return None
    if not any(marker in normalized_query for marker in ("ortalama", "gno", "agno", "gano")):
        return None
    term_gpas = _extract_term_gpa_rows(document.text)
    target = _extract_term_gpa(normalized_query)
    if term_gpas:
        matching = [row for row in term_gpas if target is None or abs(row[1] - target) < 0.005]
        if not matching:
            if target is not None:
                return f"Transkriptte {_format_number(target)} dönem ortalamasıyla bitmiş dönem tespit edemedim."
            return "Transkriptte dönem ortalaması satırları var; ancak sorudaki hedef ortalamayı güvenilir biçimde ayırt edemedim."
        lines = [f"Transkriptte {_format_number(target) if target is not None else 'istenen'} ortalamaya uyan dönemler:"]
        for label, gpa in matching[:20]:
            lines.append(f"- {label}: {_format_number(gpa)}")
        if len(matching) > 20:
            lines.append(f"- ... {len(matching) - 20} dönem daha var.")
        return "\n".join(lines)
    return (
        "Bu transkriptte dersler, toplam AKTS/kredi ve genel ortalama okunabiliyor; "
        "ancak dönem/yıl bazlı dönem ortalamalarını güvenilir biçimde ayrıştıramadım. "
        "Bu yüzden hangi dönemleri 4,00 ortalama ile bitirdiğinizi belgeye göre net söyleyemiyorum."
    )


def _extract_term_gpa_rows(text: str) -> list[tuple[str, float]]:
    rows: list[tuple[str, float]] = []
    for raw_line in text.splitlines():
        line = " ".join(raw_line.split()).strip()
        normalized = normalize_text(line)
        if not any(marker in normalized for marker in ("donem", "yariyil", "yarıyıl")):
            continue
        if not any(marker in normalized for marker in ("ortalama", "gno", "agno", "gano")):
            continue
        value = _extract_term_gpa(normalized)
        if value is None:
            continue
        label = line[:100]
        rows.append((label, value))
    return rows


def _format_schedule_document_summary(document: TranscriptDocument) -> str:
    facts = _document_facts(document)
    title_text = f" ({facts.title})" if facts.title else ""
    if not document.schedule_slots:
        return (
            f"Bu belge bir ders programi belgesi olarak gorunuyor{title_text}. Ancak PDF metin/tablo katmani "
            "guvenilir bicimde okunamadigi icin belge icindeki gun-saat-sinif satirlarini dogrudan ayristiramadim. "
            "Bu durumda ders/saat sorulari icin sistemdeki yapilandirilmis ders programi kaydina bakmak gerekir."
        )
    if document.schedule_slots:
        groups = sorted({str(slot.get("schedule_group") or "").strip() for slot in document.schedule_slots if slot.get("schedule_group")})
        days = sorted({str(slot.get("day_of_week") or "").strip() for slot in document.schedule_slots if slot.get("day_of_week")})
        return (
            f"Bu belge bir ders programı belgesi olarak görünüyor. "
            f"Belgeden {len(document.schedule_slots)} ders programı satırı ayrıştırıldı."
            + (f" Sınıf/grup: {', '.join(groups[:6])}." if groups else "")
            + (f" Günler: {', '.join(days[:7])}." if days else "")
        )
    return (
        "Bu belge bir ders programı belgesi olarak görünüyor. Ancak PDF metin/tablo katmanı güvenilir biçimde okunamadığı için "
        "belge içindeki gün-saat-sınıf satırlarını doğrudan ayrıştıramadım."
    )


def _looks_like_schedule_detail_query(normalized: str) -> bool:
    return any(
        marker in normalized
        for marker in ("ders", "program", "sinif", "saat", "gun", "cuma", "sali", "pazartesi", "carsamba", "persembe")
    )


def _answer_schedule_slots_question(*, slots: tuple[dict[str, Any], ...], query: str) -> str:
    normalized = normalize_text(query)
    filtered = list(slots)
    requested_group = _requested_schedule_group(normalized)
    requested_day = _requested_schedule_day(normalized)
    if requested_group:
        filtered = [slot for slot in filtered if _schedule_group_matches(slot.get("schedule_group"), requested_group)]
    if requested_day:
        filtered = [slot for slot in filtered if normalize_text(slot.get("day_of_week") or "") == requested_day]

    if requested_day and any(marker in normalized for marker in ("var mi", "gorunuyor mu", "ders var")):
        if not filtered:
            group_text = f" {requested_group}. sınıf" if requested_group else ""
            return f"Belgeye göre{group_text} için {_display_schedule_day(requested_day)} günü ders görünmüyor."

    if not filtered:
        parts = []
        if requested_group:
            parts.append(f"{requested_group}. sınıf")
        if requested_day:
            parts.append(_display_schedule_day(requested_day))
        scope = " / ".join(parts) if parts else "bu filtre"
        return f"Belgeden ayrıştırılan ders programı satırlarında {scope} için ders bulunamadı."

    heading_parts = []
    if requested_group:
        heading_parts.append(f"{requested_group}. sınıf")
    if requested_day:
        heading_parts.append(_display_schedule_day(requested_day))
    heading = " ".join(heading_parts) or "Belgeden ayrıştırılan ders programı"
    lines = [f"{heading} için görünen dersler:"]
    for slot in sorted(filtered, key=_schedule_slot_sort_key)[:40]:
        day = str(slot.get("day_of_week") or "")
        start = _format_schedule_time(slot.get("start_time"))
        end = _format_schedule_time(slot.get("end_time"))
        course = str(slot.get("course_name") or slot.get("course_key") or "").strip()
        room = str(slot.get("classroom") or "").strip()
        group = str(slot.get("schedule_group") or "").strip()
        time_text = f"{start}-{end}" if start and end else start or ""
        prefix = " ".join(part for part in (day, time_text) if part)
        details = " | ".join(part for part in (course, f"Derslik: {room}" if room else "", group) if part)
        lines.append(f"- {prefix} | {details}".rstrip(" |"))
    if len(filtered) > 40:
        lines.append(f"- ... {len(filtered) - 40} satır daha var.")
    return "\n".join(lines)


def _requested_schedule_group(normalized: str) -> str | None:
    match = re.search(r"\b([1-6])\s*\.?\s*sinif", normalized)
    if match:
        return match.group(1)
    return None


def _requested_schedule_day(normalized: str) -> str | None:
    day_map = {
        "pazartesi": "pazartesi",
        "sali": "sali",
        "carsamba": "carsamba",
        "persembe": "persembe",
        "cuma": "cuma",
        "cumartesi": "cumartesi",
        "pazar": "pazar",
    }
    for marker, day in day_map.items():
        if marker in normalized:
            return day
    return None


def _schedule_group_matches(value: object, requested_group: str) -> bool:
    normalized = normalize_text(value or "")
    return bool(re.search(rf"\b{re.escape(requested_group)}\s*\.?\s*sinif\b", normalized) or normalized.startswith(requested_group))


def _display_schedule_day(normalized_day: str) -> str:
    return {
        "pazartesi": "Pazartesi",
        "sali": "Salı",
        "carsamba": "Çarşamba",
        "persembe": "Perşembe",
        "cuma": "Cuma",
        "cumartesi": "Cumartesi",
        "pazar": "Pazar",
    }.get(normalized_day, normalized_day)


def _format_schedule_time(value: object) -> str:
    if value is None:
        return ""
    if hasattr(value, "strftime"):
        try:
            return value.strftime("%H:%M")
        except Exception:
            return str(value)
    return str(value)


def _schedule_slot_sort_key(slot: dict[str, Any]) -> tuple[int, str, str, str]:
    day_order = {"pazartesi": 1, "sali": 2, "carsamba": 3, "persembe": 4, "cuma": 5, "cumartesi": 6, "pazar": 7}
    day = normalize_text(slot.get("day_of_week") or "")
    return (
        day_order.get(day, 99),
        _format_schedule_time(slot.get("start_time")),
        str(slot.get("schedule_group") or ""),
        str(slot.get("course_name") or ""),
    )


def _answer_generic_document_semantic_question(*, document: TranscriptDocument, query: str) -> str | None:
    normalized_query = normalize_text(query)
    normalized_text = normalize_text(document.text)
    facts = _document_facts(document)
    if _asks_empty_document_fields(normalized_query):
        return _format_document_empty_fields(document)
    if _asks_document_field_inventory(normalized_query):
        return _format_document_field_inventory(
            document=document,
            include_values=contains_any_normalized(
                normalized_query,
                ("deger", "degeri", "degerleri", "ne", "ney", "yaziyor", "geciyor"),
            ),
        )
    if "adli sicil" in normalized_query and any(marker in normalized_query for marker in ("var", "kaydi")):
        if facts.status == "adli_sicil_ve_arsiv_kaydi_yok":
            return "Belgeye gore kisinin adli sicil kaydi ve adli sicil arsiv kaydi yoktur."
        if facts.status == "adli_sicil_kaydi_yok":
            return "Belgeye gore kisinin adli sicil kaydi yoktur."
        if facts.status == "adli_sicil_kaydi_var":
            return "Belgeye gore kisinin adli sicil kaydi vardir."
        has_no_record = "adli sicil kaydi yoktur" in normalized_text
        has_no_archive = "adli sicil arsiv kaydi yoktur" in normalized_text
        if has_no_record and has_no_archive:
            return "Belgeye göre kişinin adli sicil kaydı ve adli sicil arşiv kaydı yoktur."
        if has_no_record:
            return "Belgeye göre kişinin adli sicil kaydı yoktur."
        if "adli sicil kaydi" in normalized_text:
            return "Belgede adli sicil kaydı sonucuna ilişkin bilgi var; sonucu güvenilir biçimde özetlemek için belge metnini kontrol etmek gerekir."
    return None


def _asks_empty_document_fields(normalized_query: str) -> bool:
    return contains_any_normalized(
        normalized_query,
        (
            "digerleri bos",
            "digerleri bos mu",
            "bos olanlar",
            "hangi alanlar bos",
            "bos alanlar",
        ),
    )


def _answer_gpa_interpretation(*, document: TranscriptDocument, normalized_query: str) -> str | None:
    if document.gpa is None:
        return None
    asks_scenario = any(
        marker in normalized_query
        for marker in (
            "bu donem",
            "donem",
            "ortalama yaparsam",
            "4 ortalama",
            "kac olur",
            "ne olur",
            "hesapla",
        )
    )
    asks_improvement = any(
        marker in normalized_query
        for marker in (
            "arttir",
            "yukselt",
            "nasil arttir",
            "nasil yuksel",
            "yukseltebilir",
            "arttirabilir",
        )
    )
    if not asks_scenario and not asks_improvement:
        return None

    current = document.gpa
    weight = document.total_credit or document.total_akts
    weight_label = "kredi" if document.total_credit is not None else "AKTS"
    target_term_gpa = _extract_term_gpa(normalized_query)

    if asks_scenario:
        if weight is None:
            return (
                f"Transkriptte mevcut genel ortalamanız {_format_number(current)} görünüyor. "
                "Yeni genel ortalamayı hesaplamak için bu dönem kaç kredi/AKTS ders aldığınızı da bilmem gerekir. "
                "Formül: (mevcut GNO x mevcut toplam kredi + dönem ortalaması x dönem kredisi) / "
                "(mevcut toplam kredi + dönem kredisi)."
            )
        term_gpa_text = _format_number(target_term_gpa) if target_term_gpa is not None else "dönem ortalamanız"
        example = ""
        if target_term_gpa is not None:
            example_weight = 30.0
            projected = ((current * weight) + (target_term_gpa * example_weight)) / (weight + example_weight)
            example = (
                f" Örneğin bu dönem 30 {weight_label} üzerinden {_format_number(target_term_gpa)} ortalama yaparsanız "
                f"yaklaşık {_format_number(projected)} olur."
            )
        improvement_note = ""
        if asks_improvement and target_term_gpa is not None and target_term_gpa > current:
            improvement_note = " Bu dönem ortalamanız mevcut GNO'nuzdan yüksek olduğu için genel ortalamanız yükselir."
        return (
            f"Transkriptte mevcut genel ortalamanız {_format_number(current)} ve toplam {_format_number(weight)} {weight_label} görünüyor. "
            f"Kesin hesap için bu dönem kaç {weight_label} aldığınızı bilmem gerekir. "
            f"Formül: ({_format_number(current)} x {_format_number(weight)} + {term_gpa_text} x dönem {weight_label}) / "
            f"({_format_number(weight)} + dönem {weight_label}).{example}{improvement_note}"
        )

    if asks_improvement:
        if target_term_gpa is not None and target_term_gpa > current:
            return (
                f"Evet, bu dönem {_format_number(target_term_gpa)} ortalama yaparsanız mevcut "
                f"{_format_number(current)} GNO'nuz yükselir. Ne kadar yükseleceği bu dönem aldığınız kredi/AKTS yüküne bağlıdır."
            )
        return (
            f"Mevcut GNO'nuz {_format_number(current)}. Ortalamayı artırmak için kalan yüksek kredi/AKTS'li derslerde "
            "mevcut GNO'nuzdan daha yüksek not ortalaması yakalamanız gerekir. Yönetmelik izin veriyorsa düşük notlu "
            "dersleri yükseltmek ve başarısız görünen dersleri tamamlamak en çok etki eden adımlardır."
        )
    return None


def _extract_term_gpa(normalized_query: str) -> float | None:
    match = re.search(r"\b([0-4](?:[,.]\d{1,2})?)\s*(?:ortalama|gno|agno|gano)\b", normalized_query)
    if not match:
        match = re.search(r"\b(?:ortalama|gno|agno|gano)\s*([0-4](?:[,.]\d{1,2})?)\b", normalized_query)
    if not match:
        return None
    try:
        value = float(match.group(1).replace(",", "."))
    except ValueError:
        return None
    if 0 <= value <= 4:
        return value
    return None


def _asks_document_field_inventory(normalized_query: str) -> bool:
    return contains_any_normalized(
        normalized_query,
        (
            "hangi alanlar",
            "alanlar var",
            "alanlari ne",
            "alanlari nelerdir",
            "belgede hangi bilgiler",
            "hangi bilgiler var",
        ),
    )


def _format_document_field_inventory(*, document: TranscriptDocument, include_values: bool) -> str:
    fields = _document_facts(document).fields
    if not fields:
        return "Belgede guvenilir bicimde ayristirilmis alan tespit edemedim."
    lines = ["Belgede tespit edilen alanlar:"]
    for field in fields[:40]:
        if include_values:
            value = field.value.strip() if field.state == "filled" and field.value.strip() else "boş"
            lines.append(f"- {field.label}: {value}")
        else:
            lines.append(f"- {field.label}")
    if len(fields) > 40:
        lines.append(f"- ... {len(fields) - 40} alan daha var.")
    return "\n".join(lines)


def _format_document_empty_fields(document: TranscriptDocument) -> str:
    fields = _document_facts(document).fields
    if not fields:
        return "Belgede güvenilir biçimde ayrıştırılmış alan tespit edemedim."
    empty_fields = tuple(field for field in fields if field.state != "filled" or not field.value.strip())
    if not empty_fields:
        return "Belgede ayrıştırılan alanların tamamı dolu görünüyor."
    lines = ["Belgede boş ya da güvenilir biçimde okunamayan alanlar:"]
    for field in empty_fields[:40]:
        lines.append(f"- {field.label}")
    if len(empty_fields) > 40:
        lines.append(f"- ... {len(empty_fields) - 40} alan daha var.")
    return "\n".join(lines)


def _find_document_field(document: TranscriptDocument, query: str) -> DocumentField | None:
    match = _find_document_field_match(document, query)
    return match.field if match is not None else None


def _find_document_field_match(document: TranscriptDocument, query: str) -> _DocumentFieldMatch | None:
    fields = _document_facts(document).fields
    if not fields:
        return None
    normalized_query = normalize_text(query)
    query_tokens = _field_query_tokens(normalized_query)
    best: _DocumentFieldMatch | None = None
    for field in fields:
        field_tokens = set(field.key.split())
        if not field_tokens:
            continue
        if field.key in normalized_query:
            score = 1.0
            if len(field_tokens) == 1 and len(query_tokens) > 1:
                score = 0.7
        else:
            overlap = len(query_tokens & field_tokens)
            fuzzy_score = _field_fuzzy_score(field.key, normalized_query)
            if overlap == 0 and fuzzy_score < 0.72:
                continue
            score = max(fuzzy_score, overlap / max(min(len(field_tokens), max(len(query_tokens), 1)), 1))
            if len(field_tokens) == 1 and overlap == 1:
                score = max(score, 0.86)
        if score >= 0.72 and (best is None or score > best.score):
            best = _DocumentFieldMatch(field=field, score=score)
    return best


def _field_query_tokens(normalized_query: str) -> set[str]:
    tokens = set(_FIELD_KEY_CLEAN_RE.sub(" ", normalized_query).split())
    return tokens - {
        "bu",
        "belge",
        "belgede",
        "belgedeki",
        "belgenin",
        "ne",
        "nedir",
        "neymis",
        "hangi",
        "olarak",
        "geciyor",
        "yaziyor",
        "var",
        "mi",
        "midir",
    }


def _field_fuzzy_score(field_key: str, normalized_query: str) -> float:
    query = _field_query_label(normalized_query)
    if not query:
        return 0.0
    candidates = [field_key]
    for canonical, aliases in _FIELD_LABEL_ALIASES:
        if field_key == canonical:
            candidates.extend(aliases)
            for alias in aliases:
                if alias and alias in query:
                    return 0.96
    return max(SequenceMatcher(None, candidate, query).ratio() for candidate in candidates)


def _field_query_label(normalized_query: str) -> str:
    tokens = _FIELD_KEY_CLEAN_RE.sub(" ", normalized_query).split()
    stop_words = {
        "bu",
        "belge",
        "belgede",
        "belgedeki",
        "belgenin",
        "ne",
        "nedir",
        "neymis",
        "hangi",
        "olarak",
        "geciyor",
        "yaziyor",
        "var",
        "mi",
        "midir",
    }
    return " ".join(token for token in tokens if token not in stop_words)


async def _map_document_field_with_llm(
    *,
    llm_service: LLMGenerateProtocol,
    document: TranscriptDocument,
    query: str,
    llm_profile: str | None,
) -> DocumentField | None:
    fields = _candidate_fields_for_llm_mapping(document)
    if not fields:
        return None
    prompt = _build_llm_field_mapper_prompt(query=query, fields=fields)
    try:
        raw = await llm_service.generate(
            prompt,
            system=_FIELD_MAPPER_SYSTEM_PROMPT,
            json_mode=True,
            model_role="classifier",
            llm_profile=llm_profile,
        )
    except Exception:
        return None
    data = _parse_json_object(raw)
    if not data:
        return None
    confidence = _safe_float(data.get("confidence"))
    if confidence is None or confidence < 0.72:
        return None
    field_key = str(data.get("field_key") or "").strip()
    if not field_key or field_key.lower() in {"none", "null", "unknown"}:
        return None
    return _field_by_mapper_key(document=document, key=field_key)


def _candidate_fields_for_llm_mapping(document: TranscriptDocument) -> list[dict[str, object]]:
    fields: list[dict[str, object]] = []
    for index, field in enumerate(_document_facts(document).fields[:80]):
        if not field.key:
            continue
        fields.append(
            {
                "field_key": field.key,
                "index": index,
                "label": field.label,
                "state": field.state,
                "has_value": bool(field.value.strip()),
                "source": field.source,
            }
        )
    return fields


def _field_by_mapper_key(*, document: TranscriptDocument, key: str) -> DocumentField | None:
    normalized_key = _field_key(key)
    fields = _document_facts(document).fields
    for field in fields:
        if field.key == key or field.key == normalized_key:
            return field
    if key.isdigit():
        index = int(key)
        if 0 <= index < len(fields):
            return fields[index]
    return None


def _build_llm_field_mapper_prompt(*, query: str, fields: list[dict[str, object]]) -> str:
    return (
        f"Kullanici sorusu: {query}\n"
        f"Belgedeki alan adaylari: {json.dumps(fields, ensure_ascii=False)}\n"
        "Gorev: Kullanici hangi alanin degerini soruyor? Sadece listedeki bir field_key sec.\n"
        "Birden fazla benzer alan varsa ogrenci/kisi alanini, soru isyeri/kurum/sirket diyorsa isyeri alanini sec.\n"
        "Emin degilsen field_key null ve dusuk confidence dondur.\n"
        'Cikti JSON: {"field_key": "...", "confidence": 0.0, "reason": "..."}'
    )


def _parse_json_object(raw: str | None) -> dict[str, object] | None:
    if not raw:
        return None
    text = raw.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE).strip()
        text = re.sub(r"\s*```$", "", text).strip()
    try:
        data = json.loads(text)
    except Exception:
        match = re.search(r"\{.*\}", text, flags=re.DOTALL)
        if not match:
            return None
        try:
            data = json.loads(match.group(0))
        except Exception:
            return None
    return data if isinstance(data, dict) else None


def _safe_float(value: object) -> float | None:
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _coerce_document_intent(value: DocumentQuestionIntent | str | None) -> DocumentQuestionIntent | None:
    if value is None:
        return None
    if isinstance(value, DocumentQuestionIntent):
        return value
    try:
        return DocumentQuestionIntent(str(value))
    except ValueError:
        return None


def _requested_field_label(query: str) -> str | None:
    normalized = normalize_text(query)
    if _is_open_generic_document_question(normalized):
        return None
    if not _looks_like_field_lookup(normalized):
        return None
    cleaned = normalized
    for marker in ("belgede", "belgedeki", "belgenin", "bu belge", "bu belgede", "belgeye gore"):
        cleaned = cleaned.replace(marker, " ")
    cleaned = re.sub(r"\b(ne olarak geciyor|ne yaziyor|neymis|nedir|hangi|yaziyor|geciyor|var mi|ne)\b", " ", cleaned)
    cleaned = re.sub(r"[^a-z0-9 ]+", " ", cleaned)
    cleaned = " ".join(cleaned.split())
    if not cleaned or _is_open_generic_document_question(cleaned):
        return None
    if any(marker in cleaned for marker in ("ise yariyor", "yarar", "ediliyor", "yapilir", "basvur", "teslim", "doldur")):
        return None
    return cleaned[:80] if cleaned else None


def _is_open_generic_document_question(normalized: str) -> bool:
    if not normalized:
        return False
    open_markers = (
        "ne hakkinda",
        "neler var",
        "ne goruyorsun",
        "ne gormuyorsun",
        "bu belge ne",
        "belge ne",
        "nedir bu",
        "bu nedir",
        "ne ise yariyor",
        "ise yariyor",
        "neye yarar",
        "ne icin",
        "amaci",
        "amac",
        "ozet",
        "ozetle",
        "yorumla",
        "yorumlar misin",
        "incele",
        "inceleyip",
        "ne diyor",
        "nasil",
        "nereye",
        "nereden",
        "teslim",
        "basvur",
        "doldur",
        "imza",
        "imzalat",
        "ne yap",
    )
    return any(marker in normalized for marker in open_markers)


def _looks_like_field_lookup(normalized: str) -> bool:
    if not normalized:
        return False
    if re.search(r"\b(ne olarak geciyor|ne yaziyor|neymis|nedir|hangi|yaziyor|geciyor|var mi)\b", normalized):
        return True
    return re.search(r"\bne\b", normalized) is not None

def _format_course_list(*, title: str, courses: tuple[TranscriptCourse, ...], empty: str) -> str:
    if not courses:
        return empty
    lines = [f"{title}:"]
    for course in courses[:30]:
        details = []
        if course.akts is not None:
            details.append(f"AKTS {_format_number(course.akts)}")
        if course.credit is not None:
            details.append(f"kredi {_format_number(course.credit)}")
        if course.grade:
            details.append(f"not {course.grade}")
        suffix = f" ({', '.join(details)})" if details else ""
        lines.append(f"- {course.code} {course.name}{suffix}")
    if len(courses) > 30:
        lines.append(f"- ... {len(courses) - 30} ders daha var.")
    return "\n".join(lines)


def _format_number(value: float) -> str:
    if value.is_integer():
        return str(int(value))
    return f"{value:.2f}".replace(".", ",").rstrip("0").rstrip(",")


def _compact_raw_text(text: str, *, limit: int = 12000) -> str:
    cleaned = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if len(cleaned) <= limit:
        return cleaned
    return cleaned[:limit].rstrip()


def _build_llm_transcript_prompt(*, document: TranscriptDocument, query: str) -> str:
    rows = [
        {
            "code": course.code,
            "name": course.name,
            "credit": course.credit,
            "akts": course.akts,
            "grade": course.grade,
            "status": course.status,
            "term": course.term,
        }
        for course in document.courses[:80]
    ]
    return (
        f"Kullanıcı sorusu: {query}\n"
        "Transkript özet verisi:\n"
        f"- dosya: {document.filename}\n"
        f"- program: {document.program_name}\n"
        f"- toplam_akts: {document.total_akts}\n"
        f"- toplam_kredi: {document.total_credit}\n"
        f"- gno_agno: {document.gpa}\n"
        f"- dersler: {rows}\n"
        "Yanıtı yalnızca bu transkript verisine dayandır."
    )


def _build_llm_document_prompt(
    *,
    document: TranscriptDocument,
    query: str,
    document_intent: DocumentQuestionIntent | None = None,
) -> str:
    facts = _facts_payload(document)
    return (
        f"Kullan?c? sorusu: {query}\n"
        f"Soru tipi: {document_intent.value if document_intent else 'document_summary'}\n"
        "Y?klenen belge metni ve cikarilmis facts:\n"
        f"- dosya: {document.filename}\n"
        f"- ??kar?m_y?ntemi: {document.extraction_mode}\n"
        f"- facts: {json.dumps(facts, ensure_ascii=False)}\n"
        f"- metin:\n{document.text[:6000]}\n"
        "Yan?t? yaln?zca bu belge metnine ve cikarilmis facts alanlarina dayandir. "
        "Bos alan ile okunamayan alan ayrimini koru; belgede olmayan bilgiyi uydurma."
    )


_TRANSCRIPT_QA_SYSTEM_PROMPT = """\
Kullanıcının yüklediği transkript verisine göre kısa ve güvenli cevap ver.
Kurallar:
- Transkriptte olmayan bilgiyi uydurma.
- Program mezuniyet şartı, yönetmelik veya kurum politikası yorumu yapma; yalnız transkriptte görünen kişisel ders/AKTS/kredi/not bilgisini kullan.
- Belirsiz çıkarım gerekiyorsa bunu açıkça belirt.
- Cevabı en fazla 4 cümlede tut; ders listesi istenirse maddeleyebilirsin.
"""


_GENERIC_DOCUMENT_QA_SYSTEM_PROMPT = """\
Kullan?c?n?n y?kledi?i belgeye g?re k?sa ve g?venli cevap ver.
Kurallar:
- Belge metninde olmayan bilgiyi uydurma.
- Belge yeterli de?ilse bunu a??k?a s?yle.
- Kurum politikas? veya ki?isel hesap ??kar?m? yapma; yaln?z belgedeki metni yorumla.
- Cevab? en fazla 5 c?mlede tut; liste istenirse maddeleyebilirsin.
"""


_FIELD_MAPPER_SYSTEM_PROMPT = """\
Yuklenen belgedeki alan adaylari arasindan kullanicinin sordugu alani secen bir JSON siniflandiricisisin.
Kurallar:
- Final cevap yazma, belge degeri uretme veya tahmin etme.
- Sadece verilen alan adaylarindan field_key sec.
- Soru kisi/ogrenci hakkindaysa ogrenci/kisi alanini tercih et.
- Soru isyeri/kurum/sirket hakkindaysa isyeri alanini tercih et.
- Emin degilsen field_key null ve confidence dusuk dondur.
"""
